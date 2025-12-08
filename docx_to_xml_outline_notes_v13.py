# -*- coding: utf-8 -*-
"""
DOCX -> XML exporter (v13, style-switchable build)
Fixed structural/indentation errors and added stable table/image export.

Original features preserved:
- Heading/regex/hybrid modes
- Footnotes/endnotes extraction
- List numbering reconstruction
- Inline/paragraph style switches

Additions/repairs in this fix:
- Proper block traversal (paragraphs + tables) in document order
- Inline image extraction from runs (saved to assets/, inserted as [[IMG ...]] placeholder)
- Table serialization to [[TABLE {...}]] placeholder (cell text uses existing run-to-text logic)
- Fixed multiple indentation breaks and missing imports
"""
# ===================== CONFIG SWITCHES =====================
STYLE_FLAGS = {
    # Inline run styles
    "italic": True,
    "bold": False,
    "underline": False,
    "color": False,
    "superscript": True,
    "subscript": True,
    "tracking": False,      # letter spacing (tracking), in pt
    "font": False,          # font family
    "fontsize": False,      # font size in pt

    # Paragraph styles
    "paragraph": False,     # set to False to disable all paragraph-level attributes
}

# ===================== CODE =====================
import sys, os
import re
import json
import zipfile
import argparse
import copy
from typing import Any, Dict, List, Optional, Tuple

try:
    from regex_rules import REGEX_ORDER as DEFAULT_REGEX_ORDER, NEGATIVE_PATTERNS as DEFAULT_NEGATIVE_PATTERNS
except Exception:
    DEFAULT_REGEX_ORDER = [
        ("fixed", r'^第[\d一二三四五六七八九十百]+章[ 　\t]*'),
        ("fixed", r'^第[\d一二三四五六七八九十百]+节[ 　\t]*'),
        ("fixed", r'^第[\d一二三四五六七八九十百]+条[ 　\t]*'),
        ("numeric_dotted", None),
        ("fixed", r'^[\uFF08(]\s*\d+\s*[)\uFF09]'),
        ("fixed", r'^[\uFF08(]\s*[一二三四五六七八九十百]+\s*[)\uFF09]'),
        ("fixed", r'^\d+\s*[)）]'),
        ("fixed", r'^[一二三四五六七八九十百]+、\s*'),
    ]
    DEFAULT_NEGATIVE_PATTERNS = []

REGEX_ORDER = list(DEFAULT_REGEX_ORDER)
NEGATIVE_PATTERNS = list(DEFAULT_NEGATIVE_PATTERNS)
ACTIVE_REGEX_RULES_PATH: Optional[str] = None

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

import logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NSMAP = {"w": W_NS}
M_NSMAP = {"m": M_NS}

A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS= "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
WPG_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
V_NS = "urn:schemas-microsoft-com:vml"
NSMAP_ALL = {"w": W_NS, "a": A_NS, "r": R_NS, "wp": WP_NS, "wps": WPS_NS, "wpg": WPG_NS, "v": V_NS, "m": M_NS}
NS = NSMAP_ALL
XP_RUN_BLIPS = etree.XPath(".//w:drawing//a:blip", namespaces=NSMAP_ALL)
XP_RUN_EXTENT = etree.XPath(
    ".//w:drawing//wp:inline/wp:extent | .//w:drawing//wp:anchor/wp:extent",
    namespaces=NSMAP_ALL
)
# Text boxes (DrawingML & VML)
XP_RUN_TEXTBOX = etree.XPath(".//w:drawing//wps:txbx", namespaces=NSMAP_ALL)
XP_RUN_TEXTBOX_VML = etree.XPath(".//w:pict//v:textbox", namespaces=NSMAP_ALL)
XP_RUN_ANCHOR = etree.XPath(".//w:drawing//wp:anchor", namespaces=NSMAP_ALL)
XP_RUN_INLINE = etree.XPath(".//w:drawing//wp:inline", namespaces=NSMAP_ALL)
XP_WRAP_ANY  = etree.XPath(".//wp:anchor/*[starts-with(local-name(), 'wrap')]", namespaces=NSMAP_ALL)
XP_POS_H     = etree.XPath(".//wp:anchor/wp:positionH", namespaces=NSMAP_ALL)
XP_POS_V     = etree.XPath(".//wp:anchor/wp:positionV", namespaces=NSMAP_ALL)

# Precompiled XPaths
XP_P_OUTLINE = etree.XPath("./w:pPr/w:outlineLvl", namespaces=NSMAP)
XP_RUN_TEXTS = etree.XPath(".//w:t", namespaces=NSMAP)
XP_RUN_TABS  = etree.XPath(".//w:tab", namespaces=NSMAP)
XP_RUN_FOOTREF = etree.XPath(".//w:footnoteReference", namespaces=NSMAP)
XP_RUN_ENDREF  = etree.XPath(".//w:endnoteReference", namespaces=NSMAP)
XP_RUN_PAGEBREAK = etree.XPath(".//w:br[@w:type='page']", namespaces=NSMAP)
XP_RUN_LAST_PAGEBREAK = etree.XPath(".//w:lastRenderedPageBreak", namespaces=NSMAP)

COMPILED_FIXED: List[Optional[re.Pattern]] = []
COMPILED_NEGATIVE: List[Optional[re.Pattern]] = []
# allow “1.1标题”无空格也算编号；终止于非数字字符（可被 JSON 覆盖）
DEFAULT_NUMERIC_DOTTED_PATTERN = r'^(\d+(?:[\.．]\d+)*)(?!\d)'
NUMERIC_DOTTED_PATTERN = DEFAULT_NUMERIC_DOTTED_PATTERN
NUMERIC_DOTTED: re.Pattern = re.compile(NUMERIC_DOTTED_PATTERN)
NOTE_MARKER_TRIM = " \t\r\n()（）[]【】〔〕{}<>《》〈〉「」『』.,．、，。:：;-—﹣﹘"
NOTE_MARKER_EDGE = " \t\r\n()（）[]【】〔〕{}<>《》〈〉「」『』"
NOTE_MARKER_CHAR_SET = set(
    "0123456789０１２３４５６７８９"
    "ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫⅰⅱⅲⅳⅴⅵⅶⅷⅸⅹ"
    "一二三四五六七八九十零〇甲乙丙丁"
    "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"
)
NOTE_MARKER_SINGLE_ALPHA = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ")

def _iter_regex_config_paths(explicit_path: Optional[str] = None) -> List[str]:
    """Return possible regex config paths, ordered by priority."""
    paths: List[str] = []
    for cand in (explicit_path, os.environ.get("REGEX_RULES_PATH")):
        if cand:
            paths.append(os.path.abspath(cand))
    base_dirs = []
    if getattr(sys, "frozen", False):
        base_dirs.append(os.path.dirname(sys.executable))
    base_dirs.extend([os.getcwd(), os.path.dirname(os.path.abspath(__file__))])
    seen_dirs = set()
    for d in base_dirs:
        if not d or d in seen_dirs:
            continue
        seen_dirs.add(d)
        paths.append(os.path.join(d, "regex_rules.json"))
    out: List[str] = []
    seen_paths = set()
    for p in paths:
        if p and p not in seen_paths:
            out.append(p)
            seen_paths.add(p)
    return out

def _load_rules_from_json(path: str):
    with open(path, "r", encoding="utf-8") as fh:
        data = json.load(fh)
    order = data.get("REGEX_ORDER") or data.get("regex_order")
    negative = data.get("NEGATIVE_PATTERNS") or data.get("negative_patterns")
    numeric = data.get("NUMERIC_DOTTED") or data.get("numeric_dotted")
    return order, negative, numeric

def _emu_to_pt(val_emu: Optional[str]) -> Optional[float]:
    try:
        return float(val_emu) / 12700.0
    except Exception:
        return None

def _twip_to_pt(val: Optional[str]) -> Optional[float]:
    try:
        return float(val) / 20.0
    except Exception:
        return None

def _compile_patterns():
    global COMPILED_FIXED, COMPILED_NEGATIVE, NUMERIC_DOTTED
    COMPILED_FIXED = []
    for kind, pat in REGEX_ORDER:
        if kind == "fixed" and pat:
            COMPILED_FIXED.append(re.compile(pat))
        else:
            COMPILED_FIXED.append(None)
    COMPILED_NEGATIVE = []
    for pat in NEGATIVE_PATTERNS:
        try:
            COMPILED_NEGATIVE.append(re.compile(pat))
        except re.error:
            logger.warning(f"Invalid negative regex pattern skipped: {pat}")
            COMPILED_NEGATIVE.append(None)
    try:
        NUMERIC_DOTTED = re.compile(NUMERIC_DOTTED_PATTERN)
    except re.error:
        logger.warning(f"Invalid numeric_dotted pattern, falling back to default: {NUMERIC_DOTTED_PATTERN}")
        NUMERIC_DOTTED = re.compile(DEFAULT_NUMERIC_DOTTED_PATTERN)

def load_regex_rules(config_path: Optional[str] = None) -> Optional[str]:
    """
    Load regex rules from an external JSON file (next to the executable or via REGEX_RULES_PATH).
    Falls back to bundled defaults when no override is present.
    """
    global REGEX_ORDER, NEGATIVE_PATTERNS, ACTIVE_REGEX_RULES_PATH, NUMERIC_DOTTED_PATTERN
    chosen_path: Optional[str] = None
    for candidate in _iter_regex_config_paths(config_path):
        if not os.path.exists(candidate):
            continue
        try:
            if candidate.lower().endswith(".json"):
                order, negative, numeric = _load_rules_from_json(candidate)
            else:
                continue
            if order is None and negative is None and numeric is None:
                continue
            REGEX_ORDER = list(order or DEFAULT_REGEX_ORDER)
            NEGATIVE_PATTERNS = list(negative or DEFAULT_NEGATIVE_PATTERNS)
            NUMERIC_DOTTED_PATTERN = str(numeric) if numeric else DEFAULT_NUMERIC_DOTTED_PATTERN
            chosen_path = candidate
            # avoid leaking temp extraction path; if inside onefile temp, log generic
            onefile_temp = os.environ.get("NUITKA_ONEFILE_TEMP")
            onefile_parent = os.environ.get("NUITKA_ONEFILE_PARENT")
            def _is_under(base, path):
                try:
                    if not base:
                        return False
                    base_abs = os.path.abspath(base)
                    return os.path.commonprefix([base_abs, os.path.abspath(path)]) == base_abs
                except Exception:
                    return False
            if _is_under(onefile_temp, candidate) or _is_under(onefile_parent, candidate):
                logger.info("Regex rules loaded (bundled)")
            else:
                logger.info(f"Regex rules loaded from {candidate}")
            break
        except Exception as exc:
            logger.warning(f"Failed to load regex rules from {candidate}: {exc}")
    if chosen_path is None:
        REGEX_ORDER = list(DEFAULT_REGEX_ORDER)
        NEGATIVE_PATTERNS = list(DEFAULT_NEGATIVE_PATTERNS)
        NUMERIC_DOTTED_PATTERN = DEFAULT_NUMERIC_DOTTED_PATTERN
    ACTIVE_REGEX_RULES_PATH = chosen_path
    _compile_patterns()
    return chosen_path

load_regex_rules()

def _is_regex_excluded(text: str) -> bool:
    if not text:
        return False
    stripped = text.strip()
    if not stripped:
        return False
    for pat in COMPILED_NEGATIVE:
        if pat is not None and pat.search(stripped):
            return True
    return False

def _int_to_roman(n: int, upper: bool=True) -> str:
    if n <= 0:
        return str(n)
    vals = [
        (1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),
        (100,'C'),(90,'XC'),(50,'L'),(40,'XL'),
        (10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')
    ]
    res = []
    for v,s in vals:
        while n >= v:
            res.append(s); n -= v
    out = ''.join(res)
    return out if upper else out.lower()

def _int_to_alpha(n: int, upper: bool=False) -> str:
    if n <= 0:
        return str(n)
    s = ""
    while n > 0:
        n, r = divmod(n-1, 26)
        s = chr(ord('A' if upper else 'a') + r) + s
    return s

def _int_to_chinese(n: int) -> str:
    """
    Basic Chinese counting numerals (一, 二, ... 十, 十一, 二十, 二十一 ...)
    Enough for typical heading/list numbering.
    """
    if n <= 0:
        return str(n)
    digits = "零一二三四五六七八九"
    if n < 10:
        return digits[n]
    if n < 20:
        return "十" if n == 10 else f"十{digits[n%10]}"
    if n < 100:
        tens, ones = divmod(n, 10)
        return f"{digits[tens]}十{digits[ones] if ones else ''}"
    # 100-9999 simple handling
    units = ["", "十", "百", "千"]
    parts = []
    num = n
    unit_idx = 0
    zero_flag = False
    while num > 0 and unit_idx < len(units):
        num, rem = divmod(num, 10)
        if rem == 0:
            zero_flag = True
        else:
            if zero_flag:
                parts.append("零")
                zero_flag = False
            parts.append(units[unit_idx])
            parts.append(digits[rem])
        unit_idx += 1
    return "".join(reversed(parts)).rstrip("零")

def _strip_pstyle_marker(text: str) -> str:
    return re.sub(r'\s*\[\[PSTYLE\b[^\]]*\]\]\s*$', '', text or '')

def _parse_pstyle_marker(text: str) -> Tuple[str, Dict[str, str]]:
    """Extract [[PSTYLE ...]] from tail and return (text_wo_marker, attrs_dict)."""
    attrs: Dict[str, str] = {}
    m = re.search(r'\[\[PSTYLE\s+([^\]]+)\]\]\s*$', text or '')
    if not m:
        return text, attrs
    attr_s = m.group(1)
    # parse key="value" pairs
    for k, v in re.findall(r'([a-zA-Z\-]+)="([^"]*)"', attr_s):
        attrs[k] = v
    text_wo = text[:m.start()].rstrip()
    return text_wo, attrs

class MyDOCNode(object):
    LEVEL_TAGS = {1: "chapter", 2: "section", 3: "subsection"}

    def __init__(self, name: str, level: int, index: int = 1, parent: Optional["MyDOCNode"] = None,
                 element_type: str = "heading", properties: Optional[dict] = None):
        self.name = name or ""
        self.level = level or 0
        self.index = index or 0
        self.parent = parent
        self.children: List["MyDOCNode"] = []
        self.element_type = element_type
        self.properties = properties or {}
        self.body_paragraphs: List[str] = []  # paragraphs stored as text + optional [[PSTYLE ...]] marker
        self._pending_for_children: List[str] = []

    def add_child(self, child: "MyDOCNode"):
        self.children.append(child)

    @staticmethod
    def container_tag(level: int) -> str:
        return MyDOCNode.LEVEL_TAGS.get(level, f"level{level}")

    @staticmethod
    def heading_tag(level: int) -> str:
        return f"h{level}"

    @staticmethod
    def _escape_xml(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    @staticmethod
    def _escape_attr(val: str) -> str:
        return (val.replace("&", "&amp;")
                   .replace('"', "&quot;")
                   .replace("<", "&lt;")
                   .replace(">", "&gt;"))

    @staticmethod
    def _convert_refs_to_xml(text: str) -> str:
        # helpers
        def repl_fn(m): return f'<fnref id="{m.group(1)}"/>'
        def repl_en(m): return f'<enref id="{m.group(1)}"/>'
        def repl_i(m):  return f'<i>{m.group(1)}</i>'
        def repl_b(m):  return f'<b>{m.group(1)}</b>'
        def repl_u(m):  return f'<u>{m.group(1)}</u>'
        def repl_sup(m): return f'<sup>{m.group(1)}</sup>'
        def repl_sub(m): return f'<sub>{m.group(1)}</sub>'
        def repl_span(m):
            attrs = m.group(1) or ""
            content = m.group(2)
            out_attrs = []
            for key, data_key in (
                ("font", "data-font"),
                ("size", "data-size"),
                ("color", "data-color"),
                ("tracking", "data-tracking"),
            ):
                mm = re.search(rf'{key}="([^"]*)"', attrs)
                if mm:
                    out_attrs.append(f'{data_key}="{MyDOCNode._escape_attr(mm.group(1))}"')
            attr_out = (" " + " ".join(out_attrs)) if out_attrs else ""
            return f'<span{attr_out}>{content}</span>'

        # apply replacements
        text = re.sub(r"\[\[FNREF:(-?\d+)\]\]", repl_fn, text)
        text = re.sub(r"\[\[ENREF:(-?\d+)\]\]", repl_en, text)
        text = re.sub(r"\[\[I\]\](.*?)\[\[/I\]\]", repl_i, text, flags=re.DOTALL)
        text = re.sub(r"\[\[B\]\](.*?)\[\[/B\]\]", repl_b, text, flags=re.DOTALL)
        text = re.sub(r"\[\[U\]\](.*?)\[\[/U\]\]", repl_u, text, flags=re.DOTALL)
        text = re.sub(r"\[\[SUP\]\](.*?)\[\[/SUP\]\]", repl_sup, text, flags=re.DOTALL)
        text = re.sub(r"\[\[SUB\]\](.*?)\[\[/SUB\]\]", repl_sub, text, flags=re.DOTALL)
        text = re.sub(r'\[\[SPAN(.*?)\]\](.*?)\[\[/SPAN\]\]', repl_span, text, flags=re.DOTALL)
        return text

    @staticmethod
    def _escape_attr_dict(d: Dict[str, str]) -> Dict[str, str]:
        return {k: MyDOCNode._escape_attr(v) for k, v in d.items() if v is not None}

    def to_xml_string(self, notes: dict = None, indent: int = 0) -> str:
        pad = "  " * indent
        parts = [f'{pad}<{self.container_tag(self.level)}>']
        heading_text = self._escape_xml(self.name)
        heading_text = self._convert_refs_to_xml(heading_text)
        parts.append(f'{pad}  <{self.heading_tag(self.level)}>{heading_text}</{self.heading_tag(self.level)}>')
        if self.properties:
            parts.append(f'{pad}  <meta>')
            for k, v in self.properties.items():
                v_str = self._escape_xml(str(v))
                parts.append(f'{pad}    <prop name="{k}">{v_str}</prop>')
            parts.append(f'{pad}  </meta>')
        for para in self.body_paragraphs:
            text, pattrs = _parse_pstyle_marker(para) if STYLE_FLAGS.get("paragraph", True) else (para, {})
            escaped = self._escape_xml(text)
            escaped = self._convert_refs_to_xml(escaped)
            # build <p ...attrs>
            attrs_str = ""
            if STYLE_FLAGS.get("paragraph", True) and pattrs:
                attrs_pairs = [f'{k}="{self._escape_attr(v)}"' for k, v in pattrs.items() if v is not None and v != ""]
                if attrs_pairs:
                    attrs_str = " " + " ".join(attrs_pairs)
            parts.append(f'{pad}  <p{attrs_str}>{escaped}</p>')
        for c in self.children:
            parts.append(c.to_xml_string(notes, indent + 1))
        parts.append(f'{pad}</{self.container_tag(self.level)}>')
        return "\n".join(parts)

class Splitter:
    def matches(self, line: str) -> bool:
        raise NotImplementedError

class FixedRegexSplitter(Splitter):
    def __init__(self, pat: re.Pattern):
        self.pat = pat
    def matches(self, line: str) -> bool:
        return bool(self.pat.match(line))

class NumericDepthSplitter(Splitter):
    def __init__(self, depth: int):
        self.depth = depth
    def matches(self, line: str) -> bool:
        m = NUMERIC_DOTTED.match(line)
        if not m:
            return False
        segs = re.split(r'[\.．]', m.group(1))
        return len([s for s in segs if s]) == self.depth

class DOCXOutlineExporter:
    _NOTE_STYLE_HINTS = {"footnotereference", "endnotereference"}
    _MAX_NOTE_MARKER_LEN = 4
    def __init__(self, input_path: str, mode: str = "heading", skip_images: bool = False, skip_tables: bool = False, skip_textboxes: bool = False, regex_config_path: Optional[str] = None, regex_max_depth: Optional[int] = None, inline_list_labels: bool = True):
        assert mode in ("heading", "regex", "hybrid"), "mode must be 'heading', 'regex', or 'hybrid'"
        self.mode = mode
        self.inline_list_labels = bool(inline_list_labels)
        load_regex_rules(regex_config_path)
        self.regex_rules_path = ACTIVE_REGEX_RULES_PATH
        # None -> default 200; <=0 -> unlimited
        if regex_max_depth is None:
            self._regex_max_depth: Optional[int] = 200
        elif regex_max_depth <= 0:
            self._regex_max_depth = None
        else:
            self._regex_max_depth = int(regex_max_depth)
        self.input_path = input_path
        self.doc = Document(input_path)
        self.skip_images = bool(skip_images)
        self.skip_tables = bool(skip_tables)
        self.skip_textboxes = bool(skip_textboxes)
        self.footnotes: Dict[str, str] = {}
        self.endnotes: Dict[str, str] = {}
        self.root = MyDOCNode("root", level=0, index=0, parent=None, element_type="root")
        self._stats = {
            "body_fragments": 0,
            "table_fragments": 0,
            "image_fragments": 0,
        }
        self._last_summary: Dict[str, Any] = {}

        # Numbering (lists)
        self.num_map_abstract: Dict[int, int] = {}
        self.abstract_lvls: Dict[int, Dict[int, Dict[str, Optional[str]]]] = {}
        self.num_overrides: Dict[int, Dict[int, int]] = {}
        self.num_counters: Dict[int, List[int]] = {}
        self.assets_dir = None  # set in process()
        self.default_section_state = self._resolve_default_section_state()
        self._body_iter_items: List[Tuple[str, int, Dict[str, Any]]] = []
        self._doc_paragraphs: List[Any] = []
        self._doc_tables: List[Any] = []
        self._build_body_iter_index()
        self._frame_seq = 0
        self._word_page_width_pt, self._word_page_height_pt = self._resolve_word_page_size()
        self._word_page_seq = 1
        self._stats = {
            "body_fragments": 0,
            "table_fragments": 0,
            "image_fragments": 0,
        }
        self._last_summary: Dict[str, Any] = {}

    def _resolve_word_page_size(self) -> Tuple[float, float]:
        try:
            sect = self.doc.sections[0]
            width = getattr(sect, "page_width", None)
            height = getattr(sect, "page_height", None)
            def _to_pt(val):
                try:
                    if val is None:
                        return 0.0
                    if hasattr(val, "pt"):
                        return float(val.pt)
                    return float(val) / 20.0
                except Exception:
                    return 0.0
            wpt = _to_pt(width)
            hpt = _to_pt(height)
            return wpt, hpt
        except Exception:
            return 0.0, 0.0

    def _word_page_size_attrs(self) -> Dict[str, str]:
        attrs = {}
        if self._word_page_width_pt:
            attrs["wordPageWidth"] = f"{self._word_page_width_pt:.2f}pt"
        if self._word_page_height_pt:
            attrs["wordPageHeight"] = f"{self._word_page_height_pt:.2f}pt"
        return attrs

    @staticmethod
    def _paragraph_align_token(paragraph) -> str:
        try:
            val = paragraph.alignment
        except Exception:
            val = None
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
        }
        # optional members (not in all python-docx versions)
        for attr in ("JUSTIFY_MED", "DISTRIBUTE", "THAI_DISTRIBUTE"):
            member = getattr(WD_ALIGN_PARAGRAPH, attr, None)
            if member is not None:
                mapping[member] = "justify"
        return mapping.get(val, "")

    def _append_body_fragment(self, target_node, fragment: str):
        if fragment is None:
            return
        text = fragment.strip()
        if not text:
            return
        target_node.body_paragraphs.append(fragment)
        self._stats["body_fragments"] += 1
        if text.startswith("[[TABLE"):
            self._stats["table_fragments"] += 1

    def _count_headings(self, node: MyDOCNode) -> int:
        count = 1 if node.element_type == "heading" else 0
        for child in node.children:
            count += self._count_headings(child)
        return count

    def _collect_summary(self) -> Dict[str, Any]:
        return {
            "mode": self.mode,
            "word_paragraphs": len(self.doc.paragraphs),
            "word_tables": len(self.doc.tables),
            "image_fragments": self._stats.get("image_fragments", 0),
            "footnotes": len(self.footnotes),
            "endnotes": len(self.endnotes),
            "body_fragments": self._stats.get("body_fragments", 0),
            "table_fragments": self._stats.get("table_fragments", 0),
            "headings_detected": self._count_headings(self.root),
        }

    # ---------- Notes extraction ----------
    @staticmethod
    def _collect_text_from_p(par_el, *, skip_note_refs: bool = False) -> str:
        chunks: List[str] = []
        pending_strip_leading = False
        for node in par_el.iter():
            if node.tag == f"{{{W_NS}}}r":
                text = DOCXOutlineExporter._text_from_run(node)
                is_note_ref = False
                if skip_note_refs:
                    is_note_ref = DOCXOutlineExporter._is_note_reference_run(node, text or "")
                if is_note_ref:
                    if chunks:
                        stripped_prev = chunks[-1].rstrip(NOTE_MARKER_EDGE)
                        if stripped_prev != chunks[-1]:
                            chunks[-1] = stripped_prev
                            if not chunks[-1]:
                                chunks.pop()
                    pending_strip_leading = True
                    continue
                if not text:
                    continue
                if pending_strip_leading:
                    stripped = text.lstrip(NOTE_MARKER_EDGE)
                    if not stripped:
                        continue
                    text = stripped
                    pending_strip_leading = False
                chunks.append(text)
            elif node.tag in (f"{{{M_NS}}}oMath", f"{{{M_NS}}}oMathPara"):
                math_text = DOCXOutlineExporter._omath_to_text(node)
                if math_text:
                    chunks.append(math_text)
        if chunks:
            return "".join(chunks)
        fallback: List[str] = []
        for node in par_el.iter():
            if node.tag == f"{{{W_NS}}}t":
                fallback.append(node.text or "")
            elif node.tag == f"{{{W_NS}}}tab":
                fallback.append("\t")
            elif node.tag == f"{{{M_NS}}}t":
                fallback.append(node.text or "")
            elif node.tag == f"{{{M_NS}}}chr":
                val = node.get(f"{{{M_NS}}}val")
                if val:
                    fallback.append(val)
        return "".join(fallback)

    @staticmethod
    def _text_from_run(run_element) -> str:
        parts: List[str] = []
        for node in run_element.iter():
            if node.tag == f"{{{W_NS}}}t":
                parts.append(node.text or "")
            elif node.tag == f"{{{W_NS}}}tab":
                parts.append("\t")
        return "".join(parts)

    @staticmethod
    def _omath_to_text(math_element) -> str:
        """
        Flatten a Word OMML math node to inline text with basic sub/sup markers.
        """
        def walk(el) -> str:
            try:
                name = etree.QName(el).localname
            except Exception:
                name = el.tag.split("}", 1)[-1]

            if name == "t":
                return el.text or ""
            if name == "chr":
                return el.get(f"{{{M_NS}}}val") or ""
            if name == "r":
                return "".join(walk(ch) for ch in el)
            if name == "sSub":
                base_el = el.find("./m:e", M_NSMAP)
                sub_el = el.find("./m:sub", M_NSMAP)
                base = walk(base_el) if base_el is not None else ""
                sub = walk(sub_el) if sub_el is not None else ""
                return f"{base}[[SUB]]{sub}[[/SUB]]"
            if name == "sSup":
                base_el = el.find("./m:e", M_NSMAP)
                sup_el = el.find("./m:sup", M_NSMAP)
                base = walk(base_el) if base_el is not None else ""
                sup = walk(sup_el) if sup_el is not None else ""
                return f"{base}[[SUP]]{sup}[[/SUP]]"
            if name == "sSubSup":
                base_el = el.find("./m:e", M_NSMAP)
                sub_el = el.find("./m:sub", M_NSMAP)
                sup_el = el.find("./m:sup", M_NSMAP)
                base = walk(base_el) if base_el is not None else ""
                sub = walk(sub_el) if sub_el is not None else ""
                sup = walk(sup_el) if sup_el is not None else ""
                return f"{base}[[SUB]]{sub}[[/SUB]][[SUP]]{sup}[[/SUP]]"
            if name == "frac":
                num_el = el.find("./m:num", M_NSMAP)
                den_el = el.find("./m:den", M_NSMAP)
                num = walk(num_el) if num_el is not None else ""
                den = walk(den_el) if den_el is not None else ""
                return f"({num})/({den})"
            if name == "nary":
                sym = ""
                try:
                    chr_el = el.find("./m:naryPr/m:chr", M_NSMAP)
                    sym = chr_el.get(f"{{{M_NS}}}val") if chr_el is not None else ""
                except Exception:
                    sym = ""
                sym = sym or "∑"
                sub_el = el.find("./m:sub", M_NSMAP)
                sup_el = el.find("./m:sup", M_NSMAP)
                expr_el = el.find("./m:e", M_NSMAP)
                sub = walk(sub_el) if sub_el is not None else ""
                sup = walk(sup_el) if sup_el is not None else ""
                expr = walk(expr_el) if expr_el is not None else ""
                return f"{sym}[[SUB]]{sub}[[/SUB]][[SUP]]{sup}[[/SUP]] {expr}"
            # default: concatenate children
            return "".join(walk(ch) for ch in el) or (el.text or "")

        text = walk(math_element).strip()
        return text or "[[MATH]]"

    @classmethod
    def _is_note_reference_run(cls, run_element, text: str) -> bool:
        stripped = (text or "").strip()
        if XP_RUN_FOOTREF(run_element) or XP_RUN_ENDREF(run_element):
            return True
        rpr = run_element.find("./w:rPr", NSMAP)
        if rpr is not None:
            rstyle = rpr.find("./w:rStyle", NSMAP)
            if rstyle is not None:
                style_val = (rstyle.get(f"{{{W_NS}}}val") or "").strip().lower()
                if style_val and (style_val in cls._NOTE_STYLE_HINTS or "footnote" in style_val or "endnote" in style_val):
                    return True
            vert = rpr.find("./w:vertAlign", NSMAP)
            if vert is not None:
                vval = (vert.get(f"{{{W_NS}}}val") or "").lower()
                if vval == "superscript" and cls._looks_like_note_marker_text(stripped):
                    return True
        return cls._looks_like_note_marker_text(stripped)

    @classmethod
    def _looks_like_note_marker_text(cls, text: str) -> bool:
        if not text:
            return False
        stripped = text.strip()
        if not stripped:
            return False
        core = stripped.strip(NOTE_MARKER_TRIM)
        if not core:
            # run contains only bracket/marker characters (e.g., []、【】); treat as marker
            if stripped and all(ch in NOTE_MARKER_TRIM for ch in stripped):
                return True
            return False
        if len(core) > cls._MAX_NOTE_MARKER_LEN:
            return False
        if all(ch in NOTE_MARKER_CHAR_SET for ch in core):
            return True
        if len(core) == 1 and core in NOTE_MARKER_SINGLE_ALPHA:
            return True
        return False

    def _next_frame_id(self) -> str:
        self._frame_seq += 1
        return f"frame_{self._frame_seq:04d}"

    @staticmethod
    def _find_ancestor_anchor(node):
        cur = node
        while cur is not None:
            qname = etree.QName(cur)
            if qname.namespace == WP_NS and qname.localname in ("anchor", "inline"):
                return cur
            cur = cur.getparent()
        return None

    @staticmethod
    def _collect_anchor_attrs(anchor_el):
        data = {
            "wrap": "",
            "wrapSide": "",
            "wrapText": "",
            "posH": "",
            "posHref": "",
            "posV": "",
            "posVref": "",
            "offX": "",
            "offY": "",
            "distT": "",
            "distB": "",
            "distL": "",
            "distR": "",
            "relativeHeight": "",
            "behindDoc": "",
            "allowOverlap": "",
            "layoutInCell": "",
            "hidden": "",
            "locked": "",
            "simplePosX": "",
            "simplePosY": "",
            "effectL": "",
            "effectT": "",
            "effectR": "",
            "effectB": "",
            "sizeRelH": "",
            "sizeRelHref": "",
            "sizeRelV": "",
            "sizeRelVref": "",
            "docPrId": "",
            "docPrName": "",
            "anchorId": "",
            "anchorEditId": "",
        }
        if anchor_el is None:
            return data

        def _emu(val):
            return f"{_emu_to_pt(val):.2f}pt" if _emu_to_pt(val) is not None else ""

        data["relativeHeight"] = anchor_el.get("relativeHeight", "")
        data["behindDoc"] = anchor_el.get("behindDoc", "")
        data["allowOverlap"] = anchor_el.get("allowOverlap", "")
        data["layoutInCell"] = anchor_el.get("layoutInCell", "")
        data["hidden"] = anchor_el.get("hidden", "")
        data["locked"] = anchor_el.get("locked", "")
        WP14 = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
        data["anchorId"] = anchor_el.get(f"{{{WP14}}}anchorId", "")
        data["anchorEditId"] = anchor_el.get(f"{{{WP14}}}editId", "")

        wrap_nodes = [child for child in anchor_el if etree.QName(child).localname.startswith("wrap")]
        if wrap_nodes:
            wnode = wrap_nodes[0]
            data["wrap"] = etree.QName(wnode).localname
            data["wrapSide"] = wnode.get("wrapSide", "") or wnode.get(f"{{{WP_NS}}}wrapSide", "")
            data["wrapText"] = wnode.get("wrapText", "") or wnode.get(f"{{{WP_NS}}}wrapText", "")

        posH = anchor_el.find("wp:positionH", namespaces=NSMAP_ALL)
        if posH is not None:
            data["posHref"] = posH.get("relativeFrom", "") or posH.get(f"{{{WP_NS}}}relativeFrom", "")
            align = posH.find("wp:align", namespaces=NSMAP_ALL)
            if align is not None and align.text:
                data["posH"] = align.text
            posOffset = posH.find("wp:posOffset", namespaces=NSMAP_ALL)
            if posOffset is not None and posOffset.text:
                data["offX"] = _emu(posOffset.text)

        posV = anchor_el.find("wp:positionV", namespaces=NSMAP_ALL)
        if posV is not None:
            data["posVref"] = posV.get("relativeFrom", "") or posV.get(f"{{{WP_NS}}}relativeFrom", "")
            align = posV.find("wp:align", namespaces=NSMAP_ALL)
            if align is not None and align.text:
                data["posV"] = align.text
            posOffset = posV.find("wp:posOffset", namespaces=NSMAP_ALL)
            if posOffset is not None and posOffset.text:
                data["offY"] = _emu(posOffset.text)

        for attr, key in (("distT", "distT"), ("distB", "distB"), ("distL", "distL"), ("distR", "distR")):
            val = anchor_el.get(attr)
            if val is not None:
                data[key] = _emu(val)

        simple = anchor_el.find("wp:simplePos", namespaces=NSMAP_ALL)
        if simple is not None:
            sx = simple.get("x")
            sy = simple.get("y")
            if sx:
                data["simplePosX"] = _emu(sx)
            if sy:
                data["simplePosY"] = _emu(sy)

        effect = anchor_el.find("wp:effectExtent", namespaces=NSMAP_ALL)
        if effect is not None:
            for attr, key in (("l", "effectL"), ("t", "effectT"), ("r", "effectR"), ("b", "effectB")):
                val = effect.get(attr)
                if val is not None:
                    data[key] = _emu(val)

        size_rel_h = anchor_el.find("wp:sizeRelH", namespaces=NSMAP_ALL)
        if size_rel_h is not None:
            data["sizeRelHref"] = size_rel_h.get("relativeFrom", "")
            pct = size_rel_h.find("wp:pct", namespaces=NSMAP_ALL)
            if pct is not None and pct.text:
                try:
                    data["sizeRelH"] = f"{float(pct.text)/1000:.2f}%"
                except Exception:
                    data["sizeRelH"] = pct.text

        size_rel_v = anchor_el.find("wp:sizeRelV", namespaces=NSMAP_ALL)
        if size_rel_v is not None:
            data["sizeRelVref"] = size_rel_v.get("relativeFrom", "")
            pct = size_rel_v.find("wp:pct", namespaces=NSMAP_ALL)
            if pct is not None and pct.text:
                try:
                    data["sizeRelV"] = f"{float(pct.text)/1000:.2f}%"
                except Exception:
                    data["sizeRelV"] = pct.text

        doc_pr = anchor_el.find("wp:docPr", namespaces=NSMAP_ALL)
        if doc_pr is not None:
            data["docPrId"] = doc_pr.get("id", "")
            data["docPrName"] = doc_pr.get("name", "")

        return data

    @staticmethod
    def _collect_txbx_plain_text(txbx_node) -> str:
        lines: List[str] = []
        for p in txbx_node.findall(".//w:p", namespaces=NSMAP_ALL):
            parts: List[str] = []
            for node in p.iter():
                qname = etree.QName(node)
                if qname.namespace != W_NS:
                    continue
                if qname.localname == "t":
                    parts.append(node.text or "")
                elif qname.localname == "tab":
                    parts.append("\t")
                elif qname.localname == "br":
                    parts.append("\n")
            lines.append("".join(parts))
        return "\n".join(lines).strip("\n")

    @staticmethod
    def _pt_str_to_float(val: Optional[str]) -> Optional[float]:
        if not val:
            return None
        sval = str(val).strip()
        if sval.endswith("pt"):
            sval = sval[:-2]
        try:
            return float(sval)
        except Exception:
            return None

    def _collect_txbx_offsets(self, txbx_node):
        """
        Accumulate local offsets/sizes from nested DrawingML transforms so we can place each
        textbox relative to its group/anchor.
        """
        total_x = 0.0
        total_y = 0.0
        width = None
        height = None

        def _apply_xfrm(xfrm, take_size):
            nonlocal total_x, total_y, width, height
            if xfrm is None:
                return
            off = xfrm.find("a:off", namespaces=NSMAP_ALL)
            if off is not None:
                ox = _emu_to_pt(off.get("x"))
                oy = _emu_to_pt(off.get("y"))
                if ox is not None:
                    total_x += ox
                if oy is not None:
                    total_y += oy
            if take_size:
                ext = xfrm.find("a:ext", namespaces=NSMAP_ALL)
                if ext is not None:
                    if width is None:
                        wpt = _emu_to_pt(ext.get("cx"))
                        if wpt is not None:
                            width = wpt
                    if height is None:
                        hpt = _emu_to_pt(ext.get("cy"))
                        if hpt is not None:
                            height = hpt

        node = txbx_node.getparent()
        first_wsp_processed = False
        while node is not None:
            q = etree.QName(node)
            if q.namespace == WPS_NS and q.localname == "wsp":
                xfrm = node.find("wps:spPr/a:xfrm", namespaces=NSMAP_ALL)
                _apply_xfrm(xfrm, take_size=not first_wsp_processed)
                first_wsp_processed = True
            elif q.namespace == "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" and q.localname == "wgp":
                xfrm = node.find("wpg:grpSpPr/a:xfrm", namespaces=NSMAP_ALL)
                _apply_xfrm(xfrm, take_size=False)
            node = node.getparent()

        return total_x, total_y, width, height

    def _serialize_textbox_node(self, txbx_node) -> Optional[str]:
        anchor = self._find_ancestor_anchor(txbx_node)
        attrs = self._collect_anchor_attrs(anchor)
        local_off_x, local_off_y, local_width, local_height = self._collect_txbx_offsets(txbx_node)

        def _combine_offset(attr_name, local_val):
            base = self._pt_str_to_float(attrs.get(attr_name))
            base = base if base is not None else 0.0
            combined = base + (local_val or 0.0)
            attrs[attr_name] = f"{combined:.2f}pt"

        _combine_offset("offX", local_off_x)
        _combine_offset("offY", local_off_y)

        width = ""
        height = ""
        if local_width is not None:
            width = f"{local_width:.2f}pt"
        if local_height is not None:
            height = f"{local_height:.2f}pt"
        if not width or not height:
            if anchor is not None:
                extent = anchor.find("wp:extent", namespaces=NSMAP_ALL)
                if extent is not None:
                    cx = extent.get("cx")
                    cy = extent.get("cy")
                    val = _emu_to_pt(cx)
                    if val is not None and not width:
                        width = f"{val:.2f}pt"
                    val = _emu_to_pt(cy)
                    if val is not None and not height:
                        height = f"{val:.2f}pt"
        parent = txbx_node.getparent()
        if parent is not None:
            body_pr = parent.find("wps:bodyPr", namespaces=NSMAP_ALL)
            if body_pr is not None:
                for src, dst in (("lIns", "bodyInsetL"), ("tIns", "bodyInsetT"),
                                 ("rIns", "bodyInsetR"), ("bIns", "bodyInsetB")):
                    val = body_pr.get(src)
                    if val is not None:
                        pt = _emu_to_pt(val)
                        attrs[dst] = f"{pt:.2f}pt" if pt is not None else val
                for src, dst in (("wrap", "bodyWrap"), ("upright", "bodyUpright"),
                                 ("numCol", "bodyNumCol"), ("spcCol", "bodyColSpace"),
                                 ("rtlCol", "bodyRtlCol")):
                    val = body_pr.get(src)
                    if val is not None:
                        attrs[dst] = val
        attrs["w"] = width
        attrs["h"] = height
        attrs["id"] = self._next_frame_id()
        attrs["pageHint"] = attrs.get("anchorId") or attrs.get("docPrId") or attrs.get("docPrName") or ""
        attrs["wordPageSeq"] = str(self._word_page_seq)
        attrs.update(self._word_page_size_attrs())
        text = self._collect_txbx_plain_text(txbx_node).replace("]]", "］］")
        attr_str = " ".join(f'{k}="{MyDOCNode._escape_attr(str(v))}"' for k,v in attrs.items() if v)
        return f"[[FRAME {attr_str}]]{text}[[/FRAME]]"

    def _extract_textboxes_from_run(self, run_element) -> List[str]:
        markers: List[str] = []
        try:
            nodes = XP_RUN_TEXTBOX(run_element) or []
        except Exception:
            nodes = []
        for node in nodes:
            marker = self._serialize_textbox_node(node)
            if marker:
                markers.append(marker)
        # Legacy VML textboxes can be added here if needed
        return markers

    @staticmethod
    def _frame_payload_from_marker(marker: str) -> str:
        try:
            start = marker.index("]]") + 2
            end = marker.rindex("[[/FRAME]]")
            return marker[start:end]
        except ValueError:
            return ""

    @staticmethod
    def _is_shadow_text_after_frames(text: str, frame_fragments: List[str]) -> bool:
        """
        Detects duplicated inline text that Word keeps in the paragraph alongside anchored
        textboxes (typically the concatenation of all textbox digits). We only treat it as shadow
        text when it exactly matches the frame payloads (ignoring whitespace) or when it is purely
        whitespace/tabs.
        """
        if not text:
            return False
        if not frame_fragments:
            return False
        normalized = "".join(text.split())
        if not normalized:
            return True
        frames_joined = "".join(frame_fragments)
        frame_norm = "".join(frames_joined.split())
        if not frame_norm:
            return False
        # skip only when it is a full duplicate of textbox payloads (avoid dropping short tokens)
        return normalized == frame_norm and len(normalized) >= 3

    @staticmethod
    def _parse_notes_xml(xml_bytes: bytes) -> Dict[str, str]:
        result = {}
        if not xml_bytes:
            return result
        root = etree.fromstring(xml_bytes)
        for n in root.findall("w:footnote", NSMAP) + root.findall("w:endnote", NSMAP):
            nid = n.get(f"{{{W_NS}}}id")
            if nid is None:
                continue
            try:
                if int(nid) < 1:
                    continue
            except Exception:
                pass
            paras = n.findall(".//w:p", NSMAP)
            texts = [DOCXOutlineExporter._collect_text_from_p(p, skip_note_refs=True) for p in paras]
            result[nid] = "\n".join([t for t in texts if t.strip()])
        return result

    def extract_notes(self):
        with zipfile.ZipFile(self.input_path, "r") as z:
            foot_xml = z.read("word/footnotes.xml") if "word/footnotes.xml" in z.namelist() else None
            end_xml = z.read("word/endnotes.xml") if "word/endnotes.xml" in z.namelist() else None
            num_xml = z.read("word/numbering.xml") if "word/numbering.xml" in z.namelist() else None
        if foot_xml:
            self.footnotes = self._parse_notes_xml(foot_xml)
            logger.info(f"Footnotes parsed: {len(self.footnotes)}")
        else:
            logger.info("No footnotes.xml found.")
        if end_xml:
            self.endnotes = self._parse_notes_xml(end_xml)
            logger.info(f"Endnotes parsed: {len(self.endnotes)}")
        else:
            logger.info("No endnotes.xml found.")
        if num_xml:
            self._parse_numbering_xml(num_xml)
            logger.info("numbering.xml parsed.")
        else:
            logger.info("No numbering.xml found (list numbers may not be reconstructed).")

    # ---------- Section helpers ----------
    @staticmethod
    def _copy_section_state(state: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        if not state:
            return {}
        copied: Dict[str, Any] = {}
        for key in ("pageOrientation", "pageWidthPt", "pageHeightPt"):
            if key in state:
                copied[key] = state[key]
        margins = state.get("pageMarginsPt") if isinstance(state, dict) else None
        if isinstance(margins, dict):
            copied["pageMarginsPt"] = {k: margins[k] for k in ("top", "bottom", "left", "right") if k in margins}
        return copied

    def _build_body_iter_index(self):
        body = self.doc._element.body
        children = list(body)
        default_state = self.default_section_state
        if not children:
            self._body_iter_items = []
            self._doc_paragraphs = []
            self._doc_tables = []
            return

        para_map = {p._element: p for p in getattr(self.doc, "paragraphs", [])}
        tbl_map = {t._element: t for t in getattr(self.doc, "tables", [])}
        collected_paras: List[Any] = []
        collected_tbls: List[Any] = []
        states: List[Dict[str, Any]] = [self._copy_section_state(default_state) for _ in children]
        next_state = self._copy_section_state(default_state)
        for idx in range(len(children) - 1, -1, -1):
            child = children[idx]
            sects = [el for el in child.iter() if el.tag == qn('w:sectPr')]
            if sects:
                next_state = self._merge_section_state(default_state, sects[-1])
            states[idx] = self._copy_section_state(next_state)
        items: List[Tuple[str, int, Dict[str, Any]]] = []
        p_idx = 0
        t_idx = 0
        def append_child(child, state):
            nonlocal p_idx, t_idx, items
            tag = child.tag
            if tag.endswith('}p'):
                pobj = para_map.get(child)
                if pobj is None:
                    logger.warning("Paragraph element not found in doc.paragraphs; skipping to avoid misaligned index.")
                    return
                items.append(('p', p_idx, state))
                collected_paras.append(pobj)
                p_idx += 1
            elif tag.endswith('}tbl'):
                if not self.skip_tables:
                    tobj = tbl_map.get(child)
                    if tobj is None:
                        logger.warning("Table element not found in doc.tables; skipping to avoid misaligned index.")
                    else:
                        items.append(('tbl', t_idx, state))
                        collected_tbls.append(tobj)
                        t_idx += 1
                else:
                    t_idx += 1
            elif tag.endswith('}sdt'):
                content = child.find("./w:sdtContent", NSMAP)
                if content is not None:
                    for inner in list(content):
                        append_child(inner, state)
        for idx, child in enumerate(children):
            append_child(child, states[idx])
        self._body_iter_items = items
        self._doc_paragraphs = collected_paras
        self._doc_tables = collected_tbls

    def _table_placeholder(self, tbl_el, section_state) -> Optional[str]:
        def _unwrap_nested_table(table_el):
            """Certain word files wrap the real table inside a single-row outer table."""
            try:
                rows = table_el.findall("./w:tr", NSMAP)
            except Exception:
                return table_el
            if len(rows) != 1:
                return table_el
            tcs = rows[0].findall("./w:tc", NSMAP)
            if not tcs:
                return table_el
            candidate = None
            for tc in tcs:
                tc_text_parts = [DOCXOutlineExporter._collect_text_from_p(p) for p in tc.findall("./w:p", NSMAP)]
                # if this wrapper cell still contains visible text, keep original table
                if any((part or "").strip() for part in tc_text_parts):
                    return table_el
                inner_tbls = tc.findall(".//w:tbl", NSMAP)
                for inner in inner_tbls:
                    inner_rows = inner.findall("./w:tr", NSMAP)
                    if len(inner_rows) > 1:
                        return inner
                    if inner_rows and candidate is None:
                        candidate = inner
            return candidate or table_el

        tbl_el = _unwrap_nested_table(tbl_el)
        def _tbl_width_pt(table_el):
            tw = table_el.find("./w:tblPr/w:tblW", NSMAP)
            if tw is not None:
                t = tw.get(f"{{{W_NS}}}type")
                v = tw.get(f"{{{W_NS}}}w")
                if v and t in (None, "dxa"):
                    pt = _twip_to_pt(v)
                    if pt:
                        return max(pt, 1.0)
                if v and t == "pct":
                    try:
                        return max((float(v) / 50.0) * 4.8, 1.0)
                    except Exception:
                        pass
            return 480.0

        def _expanded_cols(row_cells):
            total = 0
            for cell in row_cells:
                try:
                    cs = int(cell.get("colspan", 1)) or 1
                except Exception:
                    cs = 1
                total += max(1, cs)
            return total

        tableWidthPt = _tbl_width_pt(tbl_el)
        headerRows = 1 if tbl_el.find("./w:tblPr/w:tblHeader", NSMAP) is not None else 0
        jc = tbl_el.find("./w:tblPr/w:jc", NSMAP)
        tableAlign = (jc.get(f"{{{W_NS}}}val") if jc is not None else None) or "left"

        borders = {"inner": 0.5, "outer": 0.75}
        tblBorders = tbl_el.find("./w:tblPr/w:tblBorders", NSMAP)

        def _edge_weight(ed):
            if ed is None:
                return None
            sz = ed.get(f"{{{W_NS}}}sz")
            try:
                return (float(sz) or 4.0) / 8.0
            except Exception:
                return None

        if tblBorders is not None:
            ins = [_edge_weight(tblBorders.find(f"./w:{n}", NSMAP)) for n in ("insideH", "insideV")]
            outs = [_edge_weight(tblBorders.find(f"./w:{n}", NSMAP)) for n in ("top", "bottom", "left", "right")]
            borders["inner"] = next((v for v in ins if v), borders["inner"])
            borders["outer"] = next((v for v in outs if v), borders["outer"])

        cellPadding = None
        tcMar = tbl_el.find("./w:tblPr/w:tblCellMar", NSMAP)
        if tcMar is not None:
            def _pad(which):
                el = tcMar.find(f"./w:{which}", NSMAP)
                return round(_twip_to_pt(el.get(f"{{{W_NS}}}w")) or 3.0, 2) if el is not None else None

            cellPadding = {"t": _pad("top"), "l": _pad("left"), "b": _pad("bottom"), "r": _pad("right")}

        rows_data: List[List[dict]] = []
        MAX_ROWS, MAX_COLS, MAX_SPAN = 500, 200, 50
        all_tr = tbl_el.findall("./w:tr", NSMAP)
        for r_idx, tr in enumerate(all_tr):
            if r_idx >= MAX_ROWS:
                break
            row_cells = []
            tcs = tr.findall("./w:tc", NSMAP)
            c_vis = 0
            for tc in tcs:
                if c_vis >= MAX_COLS:
                    break
                tcPr = tc.find("./w:tcPr", NSMAP)

                align = "left"
                p_first = tc.find("./w:p/w:pPr/w:jc", NSMAP)
                if p_first is not None and p_first.get(f"{{{W_NS}}}val"):
                    align = p_first.get(f"{{{W_NS}}}val")
                valign = "top"
                vAli = tcPr.find("./w:vAlign", NSMAP) if tcPr is not None else None
                if vAli is not None and vAli.get(f"{{{W_NS}}}val"):
                    valign = vAli.get(f"{{{W_NS}}}val")

                shading = None
                sh = tcPr.find("./w:shd", NSMAP) if tcPr is not None else None
                if sh is not None and sh.get(f"{{{W_NS}}}val") not in ("nil", "clear"):
                    shading = sh.get(f"{{{W_NS}}}fill") or sh.get(f"{{{W_NS}}}color")

                cell_text = DOCXOutlineExporter._collect_text_from_p(tc)

                gridSpan = tcPr.find("./w:gridSpan", NSMAP) if tcPr is not None else None
                colspan = 1
                if gridSpan is not None and gridSpan.get(f"{{{W_NS}}}val"):
                    try:
                        colspan = max(1, min(MAX_SPAN, int(gridSpan.get(f"{{{W_NS}}}val"))))
                    except Exception:
                        colspan = 1

                vMerge = tcPr.find("./w:vMerge", NSMAP) if tcPr is not None else None
                vm_attr = vMerge.get(f"{{{W_NS}}}val") if vMerge is not None else None
                is_continue = (vMerge is not None and vm_attr in (None, "", "continue"))
                is_restart = (vMerge is not None and vm_attr not in ("continue", "cont", "0"))

                if is_continue:
                    row_cells.append({"text": "", "colspan": 1, "rowspan": 0, "align": align, "valign": valign})
                    c_vis += 1
                    continue

                rowspan = 1
                if is_restart:
                    col_index = c_vis
                    down = r_idx + 1
                    while down < len(all_tr):
                        tlist = all_tr[down].findall("./w:tc", NSMAP)
                        if not tlist:
                            break

                        cur_col = 0
                        target_tc = None
                        for n_tc in tlist:
                            n_pr = n_tc.find("./w:tcPr", NSMAP)
                            n_grid = n_pr.find("./w:gridSpan", NSMAP) if n_pr is not None else None
                            n_cs = 1
                            if n_grid is not None and n_grid.get(f"{{{W_NS}}}val"):
                                try:
                                    n_cs = max(1, min(MAX_SPAN, int(n_grid.get(f"{{{W_NS}}}val"))))
                                except Exception:
                                    n_cs = 1
                            if cur_col <= col_index < cur_col + n_cs:
                                target_tc = n_tc
                                break
                            cur_col += n_cs

                        if target_tc is None:
                            break

                        n_pr = target_tc.find("./w:tcPr", NSMAP)
                        n_vm = n_pr.find("./w:vMerge", NSMAP) if n_pr is not None else None
                        if n_vm is None:
                            break
                        nv = n_vm.get(f"{{{W_NS}}}val")
                        if nv in (None, "", "continue", "cont", "1"):
                            rowspan += 1
                            down += 1
                            continue
                        else:
                            break

                row_cells.append({
                    "text": cell_text,
                    "colspan": colspan,
                    "rowspan": max(1, rowspan),
                    "align": align,
                    "valign": valign,
                    "shading": shading
                })
                c_vis += 1

            rows_data.append(row_cells)

        rows = len(rows_data)
        cols = 0
        for r in rows_data:
            cols = max(cols, _expanded_cols(r))

        colWidthsPt: List[float] = []
        grid_cols = tbl_el.findall("./w:tblGrid/w:gridCol", NSMAP)
        for gc in grid_cols:
            w = gc.get(f"{{{W_NS}}}w")
            pt = _twip_to_pt(w)
            if pt is not None:
                colWidthsPt.append(round(pt, 2))

        if (not colWidthsPt) or (len(colWidthsPt) != cols):
            first_tr = all_tr[0] if all_tr else None
            widths_acc = [0.0] * max(cols, len(colWidthsPt) or cols)
            used_tcW = False
            if first_tr is not None:
                tcs = first_tr.findall("./w:tc", NSMAP)
                cur_col = 0
                for tc in tcs:
                    tcPr = tc.find("./w:tcPr", NSMAP)
                    wpt = None
                    if tcPr is not None:
                        tcW = tcPr.find("./w:tcW", NSMAP)
                        if tcW is not None and tcW.get(f"{{{W_NS}}}w"):
                            wtype = tcW.get(f"{{{W_NS}}}type")
                            wval = tcW.get(f"{{{W_NS}}}w")
                            if wtype in (None, "dxa"):
                                wpt = _twip_to_pt(wval)
                            elif wtype == "pct":
                                try:
                                    pct = float(wval) / 50.0
                                    wpt = pct * 4.8
                                except Exception:
                                    wpt = None
                    gridSpan = tcPr.find("./w:gridSpan", NSMAP) if tcPr is not None else None
                    cs = 1
                    if gridSpan is not None and gridSpan.get(f"{{{W_NS}}}val"):
                        try:
                            cs = max(1, min(50, int(gridSpan.get(f"{{{W_NS}}}val"))))
                        except Exception:
                            cs = 1
                    if wpt is None:
                        wpt = tableWidthPt * (cs / float(cols or 1))
                    share = float(wpt) / float(cs or 1)
                    for k in range(cs):
                        if cur_col + k < len(widths_acc):
                            widths_acc[cur_col + k] += share
                    cur_col += cs
                    used_tcW = True
            if used_tcW:
                colWidthsPt = [round(max(1.0, v), 2) for v in widths_acc[:cols]]
            else:
                each = max(1.0, tableWidthPt / float(cols or 1))
                colWidthsPt = [round(each, 2) for _ in range(cols)]

        if len(colWidthsPt) > cols:
            colWidthsPt = colWidthsPt[:cols]
        elif len(colWidthsPt) < cols:
            missing = cols - len(colWidthsPt)
            each = max(1.0, tableWidthPt / float(cols or 1))
            colWidthsPt += [round(each, 2) for _ in range(missing)]
        s = sum(colWidthsPt) or 1.0
        scale = tableWidthPt / s
        colWidthsPt = [round(max(1.0, w * scale), 2) for w in colWidthsPt]
        total = sum(colWidthsPt) or 1.0
        colWidthFrac = [round(w / total, 6) for w in colWidthsPt]

        table_obj = {
            "rows": rows,
            "cols": cols,
            "data": rows_data,
            "colWidthsPt": colWidthsPt,
            "colWidthFrac": colWidthFrac,
            "tableWidthPt": round(tableWidthPt, 2),
            "headerRows": headerRows,
            "tableAlign": tableAlign,
        }
        if cellPadding:
            table_obj["cellPadding"] = cellPadding
        if borders:
            table_obj["borders"] = borders
        delta_state = self._section_state_delta(section_state or {}, base_state=self.default_section_state or {})
        if delta_state.get("pageOrientation"):
            table_obj["pageOrientation"] = delta_state["pageOrientation"]
        if "pageWidthPt" in delta_state:
            table_obj["pageWidthPt"] = delta_state["pageWidthPt"]
        if "pageHeightPt" in delta_state:
            table_obj["pageHeightPt"] = delta_state["pageHeightPt"]
        if "pageMarginsPt" in delta_state:
            table_obj["pageMarginsPt"] = delta_state["pageMarginsPt"]
        return "[[TABLE " + json.dumps(table_obj, ensure_ascii=False) + "]]"


    def _resolve_default_section_state(self) -> Dict[str, Any]:
        base = {
            "pageOrientation": "portrait",
            "pageWidthPt": 612.0,
            "pageHeightPt": 792.0,
            "pageMarginsPt": {"top": 72.0, "bottom": 72.0, "left": 72.0, "right": 72.0},
        }
        body_sect = getattr(self.doc._element.body, "sectPr", None)
        if body_sect is not None:
            return self._merge_section_state(base, body_sect)
        return base

    def _merge_section_state(self, base_state: Dict[str, Any], sect_pr) -> Dict[str, Any]:
        state = self._copy_section_state(base_state)
        if sect_pr is None:
            return state
        pg_sz = sect_pr.find("./w:pgSz", NSMAP)
        width_pt = state.get("pageWidthPt")
        height_pt = state.get("pageHeightPt")
        orientation = state.get("pageOrientation", "portrait")
        if pg_sz is not None:
            w_attr = pg_sz.get(f"{{{W_NS}}}w")
            h_attr = pg_sz.get(f"{{{W_NS}}}h")
            ori_attr = pg_sz.get(f"{{{W_NS}}}orient")
            w_pt = _twip_to_pt(w_attr)
            h_pt = _twip_to_pt(h_attr)
            if w_pt:
                width_pt = round(w_pt, 2)
            if h_pt:
                height_pt = round(h_pt, 2)
            if ori_attr:
                orientation = str(ori_attr).lower()
            elif w_pt and h_pt:
                orientation = "landscape" if w_pt > h_pt else "portrait"
        pg_mar = sect_pr.find("./w:pgMar", NSMAP)
        margins = self._copy_section_state(state).get("pageMarginsPt", {})
        margins = dict(margins) if margins else {}
        if pg_mar is not None:
            for key in ("top", "bottom", "left", "right"):
                val = pg_mar.get(f"{{{W_NS}}}{key}")
                if val is not None:
                    pt = _twip_to_pt(val)
                    if pt is not None:
                        margins[key] = round(pt, 2)
        state["pageOrientation"] = orientation
        if width_pt is not None:
            state["pageWidthPt"] = width_pt
        if height_pt is not None:
            state["pageHeightPt"] = height_pt
        if margins:
            state["pageMarginsPt"] = margins
        return state

    def _section_state_delta(self, state: Dict[str, Any], base_state: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        if not state:
            return {}
        default = base_state or self.default_section_state or {}
        delta: Dict[str, Any] = {}

        if state.get("pageOrientation") and state.get("pageOrientation") != default.get("pageOrientation"):
            delta["pageOrientation"] = state["pageOrientation"]

        margins = state.get("pageMarginsPt")
        def_margins = default.get("pageMarginsPt") if isinstance(default, dict) else None
        margins_delta: Dict[str, float] = {}

        def _diff(a, b, tol=0.5):
            if a is None or b is None:
                return a is not None or b is not None
            try:
                return abs(float(a) - float(b)) > tol
            except Exception:
                return a != b

        if isinstance(margins, dict):
            for key in ("top", "bottom", "left", "right"):
                if key in margins:
                    base_val = def_margins.get(key) if isinstance(def_margins, dict) else None
                    if _diff(margins[key], base_val):
                        margins_delta[key] = margins[key]
        if margins_delta:
            complete = dict(def_margins) if isinstance(def_margins, dict) else {}
            complete.update(margins_delta)
            delta["pageMarginsPt"] = complete

        return delta

    # ---------- Numbering (lists) ----------
    def _parse_numbering_xml(self, xml_bytes: bytes):
        try:
            root = etree.fromstring(xml_bytes)
        except Exception as e:
            logger.warning(f"Failed to parse numbering.xml: {e}")
            return
        for abs_num in root.findall("w:abstractNum", NSMAP):
            anid_s = abs_num.get(f"{{{W_NS}}}abstractNumId")
            if anid_s is None:
                continue
            try:
                anid = int(anid_s)
            except Exception:
                continue
            lvls: Dict[int, Dict[str, Optional[str]]] = {}
            for lvl in abs_num.findall("w:lvl", NSMAP):
                ilvl_s = lvl.get(f"{{{W_NS}}}ilvl")
                if ilvl_s is None:
                    continue
                try:
                    ilvl = int(ilvl_s)
                except Exception:
                    continue
                start = 1
                start_el = lvl.find("w:start", NSMAP)
                if start_el is not None and start_el.get(f"{{{W_NS}}}val") is not None:
                    try:
                        start = int(start_el.get(f"{{{W_NS}}}val"))
                    except Exception:
                        start = 1
                numFmt = None
                numFmt_el = lvl.find("w:numFmt", NSMAP)
                if numFmt_el is not None:
                    numFmt = numFmt_el.get(f"{{{W_NS}}}val")
                lvlText = None
                lvlText_el = lvl.find("w:lvlText", NSMAP)
                if lvlText_el is not None:
                    lvlText = lvlText_el.get(f"{{{W_NS}}}val")
                lvls[ilvl] = {"start": start, "numFmt": numFmt, "lvlText": lvlText}
            self.abstract_lvls[anid] = lvls
        for num in root.findall("w:num", NSMAP):
            numId_s = num.get(f"{{{W_NS}}}numId")
            if numId_s is None:
                continue
            try:
                numId = int(numId_s)
            except Exception:
                continue
            an_el = num.find("w:abstractNumId", NSMAP)
            if an_el is not None and an_el.get(f"{{{W_NS}}}val") is not None:
                try:
                    anid = int(an_el.get(f"{{{W_NS}}}val"))
                    self.num_map_abstract[numId] = anid
                except Exception:
                    pass
            for ov in root.findall("w:lvlOverride", NSMAP):
                ilvl_el = ov.find("w:ilvl", NSMAP)
                if ilvl_el is None or ilvl_el.get(f"{{{W_NS}}}val") is None:
                    continue
                try:
                    ilvl = int(ilvl_el.get(f"{{{W_NS}}}val"))
                except Exception:
                    continue
                start_el = ov.find(".//w:startOverride", NSMAP)
                if start_el is not None and start_el.get(f"{{{W_NS}}}val") is not None:
                    try:
                        start_val = int(start_el.get(f"{{{W_NS}}}val"))
                    except Exception:
                        start_val = None
                else:
                    start_val = None
                if start_val is not None:
                    self.num_overrides.setdefault(numId, {})[ilvl] = start_val

    def _ensure_counters(self, numId: int):
        if numId not in self.num_counters:
            counters = [0]*9
            anid = self.num_map_abstract.get(numId)
            for il in range(9):
                start = 1
                if numId in self.num_overrides and il in self.num_overrides[numId]:
                    start = self.num_overrides[numId][il]
                elif anid in self.abstract_lvls and il in self.abstract_lvls[anid]:
                    try:
                        start = int(self.abstract_lvls[anid][il].get("start") or 1)
                    except Exception:
                        start = 1
                counters[il] = start - 1
            self.num_counters[numId] = counters

    def _format_level_num(self, anid: Optional[int], ilvl: int, val: int) -> str:
        numFmt = None
        if anid in self.abstract_lvls and ilvl in self.abstract_lvls[anid]:
            numFmt = self.abstract_lvls[anid][ilvl].get("numFmt")
        if numFmt in (None, "decimal"):
            return str(val)
        if numFmt == "lowerLetter":
            return _int_to_alpha(val, upper=False)
        if numFmt == "upperLetter":
            return _int_to_alpha(val, upper=True)
        if numFmt == "lowerRoman":
            return _int_to_roman(val, upper=False)
        if numFmt == "upperRoman":
            return _int_to_roman(val, upper=True)
        if numFmt in (
            "chineseCountingThousand",
            "chineseCounting",
            "chineseLegalSimplified",
            "chineseLegal",
            "chineseCountingSimplified",
            "chineseCountingTraditional",
            "japaneseCounting",
        ):
            return _int_to_chinese(val)
        if numFmt == "bullet":
            return "•"
        return str(val)

    def _build_label_from_lvlText(self, anid: Optional[int], ilvl: int, counters: List[int]) -> str:
        lvlText = None
        if anid in self.abstract_lvls and ilvl in self.abstract_lvls[anid]:
            lvlText = self.abstract_lvls[anid][ilvl].get("lvlText")
        if not lvlText:
            cur_val = counters[ilvl]
            return self._format_level_num(anid, ilvl, cur_val) + "."
        out = lvlText
        for i in range(1, 10):
            if f"%{i}" in out:
                level_idx = i-1
                val = counters[level_idx]
                out = out.replace(f"%{i}", self._format_level_num(anid, level_idx, val))
        return out

    def list_info_for_paragraph(self, paragraph) -> Tuple[str, Dict[str, str]]:
        p_el = paragraph._element
        meta: Dict[str, str] = {}
        numPr = p_el.find("./w:pPr/w:numPr", NSMAP)

        def _numPr_from_style(style_obj):
            try:
                st = style_obj
                visited = set()
                while st and id(st) not in visited:
                    visited.add(id(st))
                    el = getattr(st, "_element", None)
                    if el is not None:
                        np = el.find("./w:pPr/w:numPr", NSMAP)
                        if np is not None:
                            return copy.deepcopy(np)
                    st = getattr(st, "base_style", None)
            except Exception:
                return None
            return None

        try:
            style_name = getattr(getattr(paragraph, "style", None), "name", None)
        except Exception:
            style_name = None

        if numPr is None:
            numPr = _numPr_from_style(getattr(paragraph, "style", None))
            if numPr is not None:
                logger.debug(f"list_info: numPr from style chain; style={style_name} text={paragraph.text!r}")
            else:
                logger.debug(f"list_info: no numPr; style={style_name} text={paragraph.text!r}")
                return "", {}

        numId_el = numPr.find("./w:numId", NSMAP)
        if numId_el is None or numId_el.get(f"{{{W_NS}}}val") is None:
            logger.debug(f"list_info: no numId; style={style_name} text={paragraph.text!r}")
            return "", {}
        ilvl_el = numPr.find("./w:ilvl", NSMAP)
        try:
            numId = int(numId_el.get(f"{{{W_NS}}}val"))
        except Exception:
            logger.debug(f"list_info: numId parse fail; style={style_name} text={paragraph.text!r}")
            return "", {}
        ilvl = 0
        if ilvl_el is not None and ilvl_el.get(f"{{{W_NS}}}val") is not None:
            try:
                ilvl = int(ilvl_el.get(f"{{{W_NS}}}val"))
            except Exception:
                ilvl = 0
        self._ensure_counters(numId)
        counters = self.num_counters[numId]
        counters[ilvl] += 1
        anid = self.num_map_abstract.get(numId)
        for j in range(ilvl+1, len(counters)):
            start = 1
            if numId in self.num_overrides and j in self.num_overrides[numId]:
                start = self.num_overrides[numId][j]
            elif anid in self.abstract_lvls and j in self.abstract_lvls[anid]:
                try:
                    start = int(self.abstract_lvls[anid][j].get("start") or 1)
                except Exception:
                    start = 1
            counters[j] = start - 1
        label = self._build_label_from_lvlText(anid, ilvl, counters)
        if label and not re.search(r'\s$', label):
            label = label + " "
        if label:
            meta["list-label"] = label
        meta["list-level"] = str(ilvl)
        numFmt = None
        if anid in self.abstract_lvls and ilvl in self.abstract_lvls[anid]:
            lvl_info = self.abstract_lvls[anid][ilvl]
            numFmt = lvl_info.get("numFmt")
            if lvl_info.get("lvlText"):
                meta["list-lvltext"] = lvl_info.get("lvlText")
        if numFmt:
            meta["list-type"] = numFmt
            meta["list-numfmt"] = numFmt
        meta["numId"] = str(numId)
        meta["abstractNumId"] = str(anid) if anid is not None else ""
        logger.debug(f"list_info: label={label!r} meta={meta} style={style_name}")
        return label, meta

    # ---------- Paragraph style helpers ----------
    @staticmethod
    def _paragraph_style_attrs(p) -> Dict[str, str]:
        if not STYLE_FLAGS.get("paragraph", True):
            return {}
        attrs: Dict[str, str] = {}
        # style name
        try:
            sname = getattr(getattr(p, "style", None), "name", None)
            if sname:
                attrs["style-name"] = str(sname)
        except Exception:
            pass
        p_el = p._element
        ppr = p_el.find("./w:pPr", NSMAP)
        # alignment
        try:
            jc = ppr.find("./w:jc", NSMAP) if ppr is not None else None
            align = jc.get(f"{{{W_NS}}}val") if jc is not None else None
            if align:
                attrs["align"] = align
        except Exception:
            pass
        # indents
        try:
            ind = ppr.find("./w:ind", NSMAP) if ppr is not None else None
            if ind is not None:
                def _pt(attr):
                    v = ind.get(f"{{{W_NS}}}{attr}")
                    if v is None:
                        return None
                    try:
                        return float(v) / 20.0
                    except Exception:
                        return None
                left = _pt("left"); right = _pt("right"); first = _pt("firstLine"); hanging = _pt("hanging")
                if left is not None: attrs["indent-left"] = f"{int(left)}pt" if float(left).is_integer() else f"{left:.1f}pt"
                if right is not None: attrs["indent-right"] = f"{int(right)}pt" if float(right).is_integer() else f"{right:.1f}pt"
                if first is not None: attrs["indent-first"] = f"{int(first)}pt" if float(first).is_integer() else f"{first:.1f}pt"
                if hanging is not None: attrs["indent-hanging"] = f"{int(hanging)}pt" if float(hanging).is_integer() else f"{hanging:.1f}pt"
        except Exception:
            pass
        # spacing
        try:
            sp = ppr.find("./w:spacing", NSMAP) if ppr is not None else None
            if sp is not None:
                def _pt(v):
                    try:
                        vv = float(v) / 20.0
                        return f"{int(vv)}pt" if float(vv).is_integer() else f"{vv:.1f}pt"
                    except Exception:
                        return None
                before = sp.get(f"{{{W_NS}}}before")
                after = sp.get(f"{{{W_NS}}}after")
                if before: attrs["space-before"] = _pt(before) or ""
                if after:  attrs["space-after"]  = _pt(after) or ""
                line = sp.get(f"{{{W_NS}}}line")
                rule = sp.get(f"{{{W_NS}}}lineRule")
                if line:
                    try:
                        line_val = float(line)
                        if rule == "auto" or rule is None:
                            # 240 = single line
                            mult = line_val / 240.0
                            attrs["line-height"] = f"{mult:.2f}x"
                            attrs["line-rule"] = "auto"
                        else:
                            # exact / atLeast: twentieths of a point
                            pt = line_val / 20.0
                            attrs["line-height"] = f"{int(pt)}pt" if float(pt).is_integer() else f"{pt:.1f}pt"
                            attrs["line-rule"] = rule
                    except Exception:
                        pass
        except Exception:
            pass
        return {k:v for k,v in attrs.items() if v}

    # ---------- Text with inline refs + numbering + styles ----------
    def _paragraph_text_with_refs(self, paragraph, include_pstyle: bool = True, include_list_prefix: bool = False) -> str:
        chunks = []
        para_align = self._paragraph_align_token(paragraph)
        # list numbering
        try:
            label, list_meta = self.list_info_for_paragraph(paragraph)
        except Exception as e:
            logger.debug(f"list label error: {e}")
            label = ""
            list_meta = {}
        want_prefix = include_list_prefix or self.inline_list_labels
        # 如果缺少有效的列表定义（无 list-type/abstractNumId），则忽略编号前缀
        has_valid_list = bool(list_meta.get("list-type") or list_meta.get("abstractNumId"))
        if label and want_prefix and has_valid_list:
            chunks.append(label)

        # Inline math (OMML) is not part of paragraph.runs; collect it so formulas are not dropped.
        math_texts: List[str] = []
        try:
            for node in paragraph._element.iter():
                if node.tag in (f"{{{M_NS}}}oMath", f"{{{M_NS}}}oMathPara"):
                    txt = self._omath_to_text(node)
                    if txt:
                        math_texts.append(txt)
        except Exception:
            pass
        if math_texts:
            chunks.extend(math_texts)

        # Traverse runs: inline images and text/notes/tabs
        para_has_frames = False
        frame_fragments: List[str] = []
        for run in paragraph.runs:
            r_el = run._element
            has_page_break = False
            try:
                if XP_RUN_LAST_PAGEBREAK(r_el):
                    has_page_break = True
                elif XP_RUN_PAGEBREAK(r_el):
                    has_page_break = True
            except Exception:
                has_page_break = False

            # ---------- IMAGES (high-fidelity [[IMG ...]] placeholder) ----------
            blips = []
            if not self.skip_images:
                try:
                    blips = XP_RUN_BLIPS(r_el) or []
                except Exception:
                    blips = []
            if blips:
                import os
                os.makedirs(self.assets_dir or ".", exist_ok=True)
                for bl in blips:
                    rid = bl.get(f"{{{R_NS}}}embed")
                    if not rid:
                        continue
                    part = getattr(run, "part", None)
                    try:
                        image_part = part.related_parts.get(rid) if part else None
                    except Exception:
                        image_part = None
                    if image_part is None:
                        continue

                    # export blob -> assets/...
                    ext = os.path.splitext(getattr(image_part, "partname", ""))[1] or ".png"
                    if not hasattr(self, "_img_seq"):
                        self._img_seq = 0
                    self._img_seq += 1
                    fname = f"image_{self._img_seq:04d}{ext}"
                    out_path = os.path.join(self.assets_dir or ".", fname)
                    try:
                        with open(out_path, "wb") as f:
                            f.write(image_part.blob)
                    except Exception:
                        pass

                    # size from wp:extent (EMU -> pt)
                    wpt = hpt = ""
                    try:
                        exts = XP_RUN_EXTENT(r_el) or []
                        if exts:
                            cx = exts[0].get("cx")
                            cy = exts[0].get("cy")
                            EMU_PER_PT = 12700.0
                            if cx:
                                wpt = f"{float(cx) / EMU_PER_PT:.2f}"
                            if cy:
                                hpt = f"{float(cy) / EMU_PER_PT:.2f}"
                    except Exception:
                        pass

                    if wpt: wpt = f"{wpt}pt"
                    if hpt: hpt = f"{hpt}pt"

                    # inline / anchor, wrap, posH/posV, offsets, distances
                    inline_el = XP_RUN_INLINE(r_el) or []
                    anchor_el = XP_RUN_ANCHOR(r_el) or []
                    anchor_node = anchor_el[0] if anchor_el else None
                    anchor_attrs = self._collect_anchor_attrs(anchor_node)
                    is_inline = bool(inline_el) and not bool(anchor_el)

                    wrap = anchor_attrs.get("wrap", "")
                    wrapSide = anchor_attrs.get("wrapSide", "")
                    wrapText = anchor_attrs.get("wrapText", "")
                    posH = anchor_attrs.get("posH", "")
                    posHref = anchor_attrs.get("posHref", "")
                    posV = anchor_attrs.get("posV", "")
                    posVref = anchor_attrs.get("posVref", "")
                    offX = anchor_attrs.get("offX", "")
                    offY = anchor_attrs.get("offY", "")
                    distT = anchor_attrs.get("distT", "") or "0pt"
                    distB = anchor_attrs.get("distB", "") or "0pt"
                    distL = anchor_attrs.get("distL", "") or "0pt"
                    distR = anchor_attrs.get("distR", "") or "0pt"
                    relativeHeight = anchor_attrs.get("relativeHeight", "")
                    behindDoc = anchor_attrs.get("behindDoc", "")
                    allowOverlap = anchor_attrs.get("allowOverlap", "")
                    layoutInCell = anchor_attrs.get("layoutInCell", "")
                    hidden = anchor_attrs.get("hidden", "")
                    locked = anchor_attrs.get("locked", "")
                    simplePosX = anchor_attrs.get("simplePosX", "")
                    simplePosY = anchor_attrs.get("simplePosY", "")
                    effectL = anchor_attrs.get("effectL", "")
                    effectT = anchor_attrs.get("effectT", "")
                    effectR = anchor_attrs.get("effectR", "")
                    effectB = anchor_attrs.get("effectB", "")
                    sizeRelH = anchor_attrs.get("sizeRelH", "")
                    sizeRelHref = anchor_attrs.get("sizeRelHref", "")
                    sizeRelV = anchor_attrs.get("sizeRelV", "")
                    sizeRelVref = anchor_attrs.get("sizeRelVref", "")
                    docPrId = anchor_attrs.get("docPrId", "")
                    docPrName = anchor_attrs.get("docPrName", "")
                    anchorId = anchor_attrs.get("anchorId", "")
                    anchorEditId = anchor_attrs.get("anchorEditId", "")
                    rotation = ""
                    flipH = ""
                    flipV = ""
                    cropT = cropB = cropL = cropR = ""
                    try:
                        xfrm = r_el.find(".//a:xfrm", NSMAP_ALL)
                        if xfrm is not None:
                            rot_val = xfrm.get("rot")
                            if rot_val:
                                rotation = f"{int(rot_val)/60000.0:.2f}"
                            flipH = xfrm.get("flipH") or ""
                            flipV = xfrm.get("flipV") or ""
                    except Exception:
                        pass
                    try:
                        crop = r_el.find(".//a:blipFill/a:srcRect", NSMAP_ALL)
                        if crop is not None:
                            def _pct(attr):
                                val = crop.get(attr)
                                if val is None:
                                    return ""
                                try:
                                    return f"{float(val)/10000.0:.2f}%"
                                except Exception:
                                    return ""
                            cropT = _pct("t")
                            cropB = _pct("b")
                            cropL = _pct("l")
                            cropR = _pct("r")
                    except Exception:
                        pass

                    # original pixel size (if available)
                    # original pixel size (true pixels, not EMU)
                    pxw = pxh = ""
                    try:
                        im = getattr(image_part, "image", None)
                        if im is not None and getattr(im, "px_width", None) and getattr(im, "px_height", None):
                            pxw = str(int(im.px_width))
                            pxh = str(int(im.px_height))
                        # fallback: if px not available, leave blank (IDML绔寜 w/h pt 澶勭悊鍗冲彲)
                    except Exception:
                        pass

                    if (not posV) and offY:
                        posV = "paragraph" 
                    try:
                        out_path = out_path.replace("\\", "/")
                    except Exception:
                        try:
                            out_path = str(out_path).replace("\\", "/")
                        except Exception:
                            out_path = ""


                    inline_flag = "1" if is_inline else "0"

                    img_tpl = (
                        '[[IMG src="{src}" w="{w}" h="{h}" pxw="{pxw}" pxh="{pxh}" '
                        'align="{align}" inline="{inline_}" wrap="{wrap}" wrapSide="{wrapSide}" wrapText="{wrapText}" '
                        'posH="{posH}" posHref="{posHref}" posV="{posV}" posVref="{posVref}" rotation="{rotation}" '
                        'flipH="{flipH}" flipV="{flipV}" offX="{offX}" offY="{offY}" distT="{distT}" distB="{distB}" '
                        'distL="{distL}" distR="{distR}" cropT="{cropT}" cropB="{cropB}" cropL="{cropL}" cropR="{cropR}" '
                        'relativeHeight="{relativeHeight}" behindDoc="{behindDoc}" allowOverlap="{allowOverlap}" '
                        'layoutInCell="{layoutInCell}" hidden="{hidden}" locked="{locked}" simplePosX="{simplePosX}" '
                        'simplePosY="{simplePosY}" effectL="{effectL}" effectT="{effectT}" effectR="{effectR}" effectB="{effectB}" '
                        'sizeRelH="{sizeRelH}" sizeRelHref="{sizeRelHref}" sizeRelV="{sizeRelV}" sizeRelVref="{sizeRelVref}" '
                        'docPrId="{docPrId}" docPrName="{docPrName}" anchorId="{anchorId}" anchorEditId="{anchorEditId}" '
                        'wordPageWidth="{wordPageWidth}" wordPageHeight="{wordPageHeight}" wordPageSeq="{wordPageSeq}"]]'
                    )
                    fmt_vals = dict(
                        src=out_path,
                        w=wpt,
                        h=hpt,
                        pxw=pxw,
                        pxh=pxh,
                        align=para_align,
                        inline_=inline_flag,
                        wrap=wrap,
                        wrapSide=wrapSide,
                        wrapText=wrapText,
                        posH=posH,
                        posHref=posHref,
                        posV=posV,
                        posVref=posVref,
                        rotation=rotation,
                        flipH=flipH,
                        flipV=flipV,
                        offX=offX,
                        offY=offY,
                        distT=distT,
                        distB=distB,
                        distL=distL,
                        distR=distR,
                        cropT=cropT,
                        cropB=cropB,
                        cropL=cropL,
                        cropR=cropR,
                        relativeHeight=relativeHeight,
                        behindDoc=behindDoc,
                        allowOverlap=allowOverlap,
                        layoutInCell=layoutInCell,
                        hidden=hidden,
                        locked=locked,
                        simplePosX=simplePosX,
                        simplePosY=simplePosY,
                        effectL=effectL,
                        effectT=effectT,
                        effectR=effectR,
                        effectB=effectB,
                        sizeRelH=sizeRelH,
                        sizeRelHref=sizeRelHref,
                        sizeRelV=sizeRelV,
                        sizeRelVref=sizeRelVref,
                        docPrId=docPrId,
                        docPrName=docPrName,
                        anchorId=anchorId,
                        anchorEditId=anchorEditId,
                        wordPageWidth="",
                        wordPageHeight="",
                        wordPageSeq="",
                    )
                    fmt_vals.update(self._word_page_size_attrs())
                    fmt_vals["wordPageSeq"] = str(self._word_page_seq)
                    chunks.append(img_tpl.format(**fmt_vals))
                    try:
                        self._stats["image_fragments"] = self._stats.get("image_fragments", 0) + 1
                    except Exception:
                        self._stats["image_fragments"] = self._stats.get("image_fragments", 0) + 1

            frame_markers = [] if self.skip_textboxes else self._extract_textboxes_from_run(r_el)
            if frame_markers:
                chunks.extend(frame_markers)
                for marker in frame_markers:
                    payload = self._frame_payload_from_marker(marker)
                    if payload:
                        frame_fragments.append(payload)
                para_has_frames = True
                continue
            # ---------- text + notes + tabs ----------
            parts_run = []
            for t in XP_RUN_TEXTS(r_el):
                if t.text:
                    parts_run.append(t.text)
            for ref in XP_RUN_FOOTREF(r_el):
                rid = ref.get(f"{{{W_NS}}}id")
                parts_run.append(f"[[FNREF:{rid}]]")
            for ref in XP_RUN_ENDREF(r_el):
                rid = ref.get(f"{{{W_NS}}}id")
                parts_run.append(f"[[ENREF:{rid}]]")
            if XP_RUN_TABS(r_el):
                parts_run.append("\t")
            run_text = "".join(parts_run)
            if run_text:
                if para_has_frames and self._is_shadow_text_after_frames(run_text, frame_fragments):
                    continue
                # --- INLINE STYLE WRAP (from STYLE_FLAGS) ---
                try:
                    # Detect run styles robustly
                    ital = getattr(run, "italic", None)
                    bold = getattr(run, "bold", None)
                    under = getattr(run, "underline", None)
                    fobj = getattr(run, "font", None)
                    if ital is None and fobj is not None:
                        ital = getattr(fobj, "italic", None)
                    if bold is None and fobj is not None:
                        bold = getattr(fobj, "bold", None)
                    if under is None and fobj is not None:
                        under = getattr(fobj, "underline", None)

                    # Fallback to raw rPr
                    try:
                        rpr = r_el.find(".//w:rPr", NSMAP)
                    except Exception:
                        rpr = None

                    def rpr_has(tag):
                        if rpr is None:
                            return False
                        el = rpr.find(tag, NSMAP)
                        if el is None:
                            return False
                        val = el.get(f"{{{W_NS}}}val")
                        return (val in (None, "", "1", "true", "True")) or (
                            tag == 'w:u' and val not in ("none", "0", "false")
                        )

                    # superscript / subscript
                    supers = getattr(fobj, "superscript", False) if fobj is not None else False
                    sub = getattr(fobj, "subscript", False) if fobj is not None else False
                    if not supers and not sub and rpr is not None:
                        va = rpr.find("w:vertAlign", NSMAP)
                        if va is not None:
                            v = va.get(f"{{{W_NS}}}val")
                            supers = (v == "superscript")
                            sub = (v == "subscript")

                    # font family / size / color
                    font_name = None
                    font_size = None
                    color_hex = None
                    tracking = None

                    try:
                        if fobj is not None and getattr(fobj, "name", None):
                            font_name = str(fobj.name)
                        elif rpr is not None:
                            rf = rpr.find("w:rFonts", NSMAP)
                            if rf is not None:
                                font_name = rf.get(f"{{{W_NS}}}ascii") or rf.get(f"{{{W_NS}}}hAnsi")
                    except Exception:
                        pass

                    try:
                        if fobj is not None and getattr(fobj, "size", None):
                            try:
                                font_size = float(fobj.size.pt)
                            except Exception:
                                font_size = None
                        elif rpr is not None:
                            sz = rpr.find("w:sz", NSMAP)
                            if sz is not None:
                                val = sz.get(f"{{{W_NS}}}val")
                                if val:
                                    font_size = float(val) / 2.0
                    except Exception:
                        pass

                    try:
                        if rpr is not None:
                            c = rpr.find("w:color", NSMAP)
                            if c is not None:
                                v = c.get(f"{{{W_NS}}}val")
                                if v and v.lower() not in ("auto",):
                                    if len(v) in (6, 3):
                                        color_hex = "#" + v if not v.startswith("#") else v
                    except Exception:
                        pass

                    # Apply wrappers based on STYLE_FLAGS
                    span_attrs = []
                    if STYLE_FLAGS.get("font", False) and font_name:
                        span_attrs.append(f'font="{font_name}"')
                    if STYLE_FLAGS.get("fontsize", False) and font_size:
                        span_attrs.append(
                            f'size="{int(font_size) if float(font_size).is_integer() else font_size}"'
                        )
                    if STYLE_FLAGS.get("color", False) and color_hex:
                        span_attrs.append(f'color="{color_hex}"')
                    if STYLE_FLAGS.get("tracking", False) and tracking:
                        span_attrs.append(f'tracking="{tracking}"')
                    if span_attrs:
                        run_text = f"[[SPAN {' '.join(span_attrs)}]]{run_text}[[/SPAN]]"

                    if STYLE_FLAGS.get("superscript", False) and supers:
                        run_text = f"[[SUP]]{run_text}[[/SUP]]"
                    elif STYLE_FLAGS.get("subscript", False) and sub:
                        run_text = f"[[SUB]]{run_text}[[/SUB]]"

                    if STYLE_FLAGS.get("underline", False):
                        is_u = (under is True) or rpr_has("w:u")
                        if is_u:
                            run_text = f"[[U]]{run_text}[[/U]]"

                    if STYLE_FLAGS.get("bold", False):
                        is_b = (bold is True) or rpr_has("w:b")
                        if is_b:
                            run_text = f"[[B]]{run_text}[[/B]]"

                    if STYLE_FLAGS.get("italic", True):
                        is_i = (ital is True) or rpr_has("w:i")
                        if is_i:
                            run_text = f"[[I]]{run_text}[[/I]]"
                except Exception:
                    pass
                chunks.append(run_text)
            if has_page_break:
                self._word_page_seq += 1

        out = "".join(chunks).strip()

        # Paragraph style marker (only for body lines, not headings)
        if include_pstyle and STYLE_FLAGS.get("paragraph", True):
            pattrs = MyDOCNode._escape_attr_dict(self._paragraph_style_attrs(paragraph))
            if pattrs:
                kv = " ".join([f'{k}="{v}"' for k, v in pattrs.items()])
                out += f' [[PSTYLE {kv}]]'
        return out

    # ---------- Heading detection (heading mode) ----------
    @staticmethod
    def _style_based_heading_level(style_name: Optional[str]) -> Optional[int]:
        if not style_name:
            return None
        m = re.match(r"Heading\s+(\d+)$", str(style_name).strip(), flags=re.IGNORECASE)
        if m:
            try:
                return int(m.group(1))
            except Exception:
                return None
        return None

    @staticmethod
    def _outline_level_from_p(paragraph) -> Optional[int]:
        el = paragraph._element
        outlvl_nodes = XP_P_OUTLINE(el)
        if outlvl_nodes:
            try:
                val = outlvl_nodes[0].get(f"{{{W_NS}}}val")
                if val is not None:
                    return int(val) + 1
            except Exception:
                return None
        return None

    # ---------- Regex mode helpers ----------
    def _regex_choose_splitter(self, lines: List[str]) -> Optional[object]:
        for line in lines:
            s = line.strip()
            if not s:
                continue
            for idx, (kind, pat) in enumerate(REGEX_ORDER):
                if kind == "fixed":
                    cp = COMPILED_FIXED[idx]
                    if cp is not None and cp.match(s):
                        if _is_regex_excluded(s):
                            continue
                        return FixedRegexSplitter(cp)
                elif kind == "numeric_dotted":
                    m = NUMERIC_DOTTED.match(s)
                    if m:
                        rest = s[m.end():].lstrip()
                        token = m.group(1) or ""
                        if rest and rest[0] in "年月日" and len(token) >= 4:
                            continue
                        if not rest:
                            continue
                        depth = len([seg for seg in re.split(r'[\.．]', token) if seg])
                        if _is_regex_excluded(s):
                            continue
                        return NumericDepthSplitter(depth)
        return None

    def _regex_split_into_nodes(self, parent: MyDOCNode, level: int, lines: List[str]) -> List[MyDOCNode]:
        nodes: List[MyDOCNode] = []
        splitter = self._regex_choose_splitter(lines)
        found_first = False
        current: Optional[MyDOCNode] = None

        if splitter is None:
            for ln in lines:
                if ln and ln.strip():
                    self._append_body_fragment(parent, ln)
            return nodes

        for ln in lines:
            s = ln.strip()
            if s.startswith("[[LAYOUT"):
                # Keep layout markers in place with the current scope
                if found_first and current is not None:
                    current._pending_for_children.append(ln)
                else:
                    self._append_body_fragment(parent, ln)
                continue
            if splitter.matches(s) and not _is_regex_excluded(s):
                index = len(nodes) + 1
                props = {"mode": "regex"}
                title = _strip_pstyle_marker(s)
                node = MyDOCNode(name=title, level=level, index=index, parent=parent,
                                 element_type="heading", properties=props)
                parent.add_child(node)
                nodes.append(node)
                current = node
                found_first = True
            else:
                if not found_first:
                    if s:
                        self._append_body_fragment(parent, ln)
                else:
                    current._pending_for_children.append(ln)
        return nodes

    def _regex_build_recursive(self, parent: MyDOCNode, level: int, lines: List[str], max_depth: Optional[int] = 200):
        if max_depth is not None and level > max_depth:
            for ln in lines:
                if ln and ln.strip():
                    self._append_body_fragment(parent, ln)
            return
        children = self._regex_split_into_nodes(parent, level, lines)
        if not children:
            return
        for child in children:
            if child._pending_for_children:
                child_lines = child._pending_for_children
                child._pending_for_children = []
                self._regex_build_recursive(child, level + 1, child_lines, max_depth=max_depth)

    def _regex_finalize_pending_to_body(self, node: MyDOCNode):
        if node._pending_for_children:
            for ln in node._pending_for_children:
                if ln and ln.strip():
                    self._append_body_fragment(node, ln)
            node._pending_for_children = []
        for ch in node.children:
            self._regex_finalize_pending_to_body(ch)

    # ---------- Build tree ----------
    def build_tree(self):
        if self.mode == "heading":
            self._build_tree_heading_mode()
        elif self.mode == "regex":
            self._build_tree_regex_mode()
        else:
            logger.info("Building hierarchy (hybrid): heading first, then regex refine bodies...")
            self._build_tree_heading_mode()
            self._refine_bodies_with_regex(self.root)
            logger.info("Hybrid build complete.")

    def _iter_block_items(self):
        """Yield ('p'/'tbl', object, section_state) in document order."""
        for kind, index, state in self._body_iter_items:
            if kind == 'p':
                if index >= len(self._doc_paragraphs):
                    logger.warning(f"Paragraph index {index} out of range (len={len(self._doc_paragraphs)}); skipped.")
                    continue
                yield ('p', self._doc_paragraphs[index], self._copy_section_state(state))
            elif kind == 'tbl':
                if self.skip_tables:
                    continue
                if index >= len(self._doc_tables):
                    logger.warning(f"Table index {index} out of range (len={len(self._doc_tables)}); skipped.")
                    continue
                yield ('tbl', self._doc_tables[index], self._copy_section_state(state))

    def _build_tree_heading_mode(self):
        logger.info("Building hierarchy (heading mode)...")
        current = self.root
        stack: List[MyDOCNode] = [self.root]

        def _tbl_width_pt(tbl_el):
            tw = tbl_el.find("./w:tblPr/w:tblW", NSMAP)
            if tw is not None:
                t = tw.get(f"{{{W_NS}}}type")
                v = tw.get(f"{{{W_NS}}}w")
                if v and t in (None, "dxa"):
                    pt = _twip_to_pt(v)
                    if pt:
                        return max(pt, 1.0)
                if v and t == "pct":
                    try:
                        return max((float(v) / 50.0) * 4.8, 1.0)  # 100%≈480pt
                    except Exception:
                        pass
            return 480.0

        def _expanded_cols(row_cells):
            n = 0
            for cell in row_cells:
                try:
                    cs = int(cell.get("colspan", 1)) or 1
                except Exception:
                    cs = 1
                n += max(1, cs)
            return n

        for kind, obj, section_state in self._iter_block_items():
            if kind == "p":
                p = obj
                raw_text = (p.text or "").strip()

                level = self._outline_level_from_p(p)
                if level is None:
                    try:
                        style_name = getattr(p.style, "name", None)
                    except Exception:
                        style_name = None
                    lvl2 = self._style_based_heading_level(style_name)
                    level = lvl2 if lvl2 is not None else 0

                if level > 0 and (raw_text or True):
                    # 调整栈
                    while len(stack) > level:
                        stack.pop()
                    while len(stack) < level:
                        dummy = MyDOCNode(name="", level=len(stack), index=1, parent=stack[-1], element_type="heading")
                        stack[-1].add_child(dummy)
                        stack.append(dummy)

                    parent = stack[level - 1]
                    index = sum(1 for c in parent.children if c.level == level) + 1
                    props = {"style": getattr(p.style, "name", None), "outline_level": level, "mode": "heading"}
                    heading_text = self._paragraph_text_with_refs(p, include_pstyle=False) or raw_text
                    node = MyDOCNode(name=heading_text, level=level, index=index,
                                     parent=parent, element_type="heading", properties=props)
                    parent.add_child(node)
                    if len(stack) == level:
                        stack.append(node)
                    else:
                        stack[level] = node
                        stack = stack[:level + 1]
                    current = node
                else:
                    text_with_refs = self._paragraph_text_with_refs(p, include_pstyle=True)
                    if text_with_refs or raw_text:
                        target = current if current is not None else self.root
                        if text_with_refs.strip():
                            self._append_body_fragment(target, text_with_refs)

            elif kind == "tbl":
                ph = self._table_placeholder(obj._element, section_state)
                if ph:
                    target = stack[-1] if stack else self.root
                    self._append_body_fragment(target, ph)
                continue


        logger.info("Heading tree build complete.")

    def _build_tree_regex_mode(self):
        logger.info("Building hierarchy (regex mode, hierarchical segmentation with dynamic numeric depth)...")
        lines: List[str] = []
        for kind, obj, section_state in self._iter_block_items():
            if kind == "layout":
                marker = self._layout_marker_from_state(section_state)
                if marker:
                    lines.append(marker)
                continue
            if kind == "p":
                line = self._paragraph_text_with_refs(obj, include_pstyle=True)
                lines.append(line)
            elif kind == "tbl":
                ph = self._table_placeholder(obj._element, section_state)
                if ph:
                    lines.append(ph)
        self._regex_build_recursive(self.root, level=1, lines=lines, max_depth=self._regex_max_depth)
        self._regex_finalize_pending_to_body(self.root)
        logger.info("Regex tree build complete.")

    # ---------- Hybrid refinement ----------
    def _refine_bodies_with_regex(self, node: MyDOCNode):
        if node.level >= 1 and node.body_paragraphs:
            lines = node.body_paragraphs[:]
            node.body_paragraphs = []
            before = len(node.children)
            self._regex_build_recursive(node, level=node.level + 1, lines=lines, max_depth=self._regex_max_depth)
            self._regex_finalize_pending_to_body(node)
            after = len(node.children)
            if after > before:
                newkids = node.children[before:]
                oldkids = node.children[:before]
                node.children = newkids + oldkids
        for ch in list(node.children):
            self._refine_bodies_with_regex(ch)

    # ---------- XML export ----------
    @staticmethod
    def _escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _notes_to_xml(self) -> str:
        parts = []
        if self.footnotes:
            parts.append("  <footnotes>")
            for nid, body in sorted(self.footnotes.items(), key=lambda kv: int(kv[0])):
                parts.append(f'    <footnote id="{nid}">{self._escape(body)}</footnote>')
            parts.append("  </footnotes>")
        if self.endnotes:
            parts.append("  <endnotes>")
            for nid, body in sorted(self.endnotes.items(), key=lambda kv: int(kv[0])):
                parts.append(f'    <endnote id="{nid}">{self._escape(body)}</endnote>')
            parts.append("  </endnotes>")
        return "\n".join(parts)

    def to_xml(self, output_path: str):
        logger.info("Writing XML")
        parts = ['<?xml version="1.0" encoding="UTF-8"?>', "<document>"]
        if self.root.body_paragraphs:
            parts.append("  <body>")
            for para in self.root.body_paragraphs:
                text, pattrs = _parse_pstyle_marker(para) if STYLE_FLAGS.get("paragraph", True) else (para, {})
                esc = MyDOCNode._escape_xml(text)
                esc = MyDOCNode._convert_refs_to_xml(esc)
                attrs_str = ""
                if STYLE_FLAGS.get("paragraph", True) and pattrs:
                    kvs = [f'{k}="{MyDOCNode._escape_attr(v)}"' for k,v in pattrs.items() if v]
                    if kvs:
                        attrs_str = " " + " ".join(kvs)
                parts.append(f"    <p{attrs_str}>{esc}</p>")
            parts.append("  </body>")
        for child in self.root.children:
            parts.append(child.to_xml_string(notes={"footnotes": self.footnotes, "endnotes": self.endnotes}, indent=1))
        notes_xml = self._notes_to_xml()
        if notes_xml:
            parts.append(notes_xml)
        parts.append("</document>")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(parts))
        logger.info("Done.")

    def _count_headings(self, node: MyDOCNode) -> int:
        count = 1 if node.element_type == "heading" else 0
        for child in node.children:
            count += self._count_headings(child)
        return count

    def _collect_summary(self) -> Dict[str, Any]:
        return {
            "mode": self.mode,
            "word_paragraphs": len(self.doc.paragraphs),
            "word_tables": len(self.doc.tables),
            "image_fragments": self._stats.get("image_fragments", 0),
            "footnotes": len(self.footnotes),
            "endnotes": len(self.endnotes),
            "body_fragments": self._stats.get("body_fragments", 0),
            "table_fragments": self._stats.get("table_fragments", 0),
            "headings_detected": self._count_headings(self.root),
        }

    def summary(self) -> Dict[str, Any]:
        return self._last_summary or self._collect_summary()

    def process(self, output_path: str):
        # set assets dir next to output
        outdir = os.path.dirname(os.path.abspath(output_path)) or "."
        self.assets_dir = os.path.join(outdir, "assets")
        self.extract_notes()
        self.build_tree()
        self.to_xml(output_path)
        self._last_summary = self._collect_summary()
        return self._last_summary
        self._last_summary = self._collect_summary()
        return self._last_summary

def main(argv=None):
    parser = argparse.ArgumentParser(description="DOCX -> XML exporter (heading/regex/hybrid) with style switches")
    parser.add_argument("--mode", choices=["heading", "regex", "hybrid"], default="heading", help="Detection mode")
    parser.add_argument("--regex-config", help="Path to regex_rules.json override (optional)")
    parser.add_argument("--regex-max-depth", type=int, help="Max depth for regex segmentation (0 for unlimited, default 200)")
    parser.add_argument("--no-inline-list-labels", dest="inline_list_labels", action="store_false", help="不把列表编号前缀写回正文，仅保留为元数据（默认写回，与 Word 视觉一致）")
    parser.set_defaults(inline_list_labels=True)
    parser.add_argument("input", help="Input .docx path")
    parser.add_argument("output", help="Output .xml path")
    args = parser.parse_args(argv)

    exporter = DOCXOutlineExporter(
        args.input,
        mode=args.mode,
        regex_config_path=args.regex_config,
        regex_max_depth=args.regex_max_depth,
        inline_list_labels=args.inline_list_labels,
    )
    exporter.process(args.output)
    # print(f"[OK] mode={args.mode} XML saved -> {args.output}")

if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
